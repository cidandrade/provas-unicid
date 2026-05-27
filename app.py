# -*- coding: utf-8 -*-
"""
Gerador de Provas Unicid - Versão Web
Versão 3.6.3
Gera provas com questões de múltipla escolha e dissertativas,
já no formato da Unicid

Autor: Prof.Me. Cid R. Andrade (profandrade@gmail.com)
Co-Autor: Prof.Me. Rafael Cotrin (a partir da v3.4.0)
Data desta versão: maio/2026

Formato do arquivo XLSX de questões OBJETIVAS:
- Coluna A: Enunciado
- Coluna B: Resposta Correta
- Colunas C–F: Distratores (4 opções incorretas)
A primeira linha deve ser a primeira questão (sem cabeçalho).

Formato do arquivo XLSX de questões DISSERTATIVAS:
- Coluna A: Enunciado
A primeira linha deve ser a primeira questão (sem cabeçalho).

Este programa é Software Livre licenciado sob a GPL v3+.
Veja https://www.gnu.org/licenses/ para mais detalhes.

ChangeLog
3.6.3 maio/2026: Geração de imagens ilustrativas via DALL-E 3 (opcional); upload
                  manual de imagens; fluxo de aprovação pelo professor antes de
                  inserir no DOCX; ajuste no nº de questões geradas por chamada IA
3.6.2 maio/2026: Uniformiza nome do autor para Prof.Me. Cid R. Andrade; corrige
                  textos do Manual/FAQ sobre PDF (LibreOffice ou Word por plataforma)
3.6.1 maio/2026: Randomização da IA: variação automática de instrução em cada
                  chamada aos prompts; pool dissertativas aumentado para 8;
                  aviso de variedade quando pool é pequeno
3.6.0 maio/2026: Integração com Claude API para importação de documentos não
                  estruturados (DOCX, XLSX, TXT/Google Forms); chave API
                  inserida pelo usuário e apagada após o processamento
3.5.2 maio/2026: Seção dissertativa usa tipo 'continuous' para balancear colunas
                  objetivas; CoInitialize/CoUninitialize para PDF em threads
3.5.1 maio/2026: Questões dissertativas passam a ocupar seção de 1 coluna
                  (largura total da página), separada da seção de 2 colunas
                  das objetivas via quebra de seção 'nextPage'; dissertativas
                  sempre iniciam em página nova com layout completo;
                  erros de conversão PDF agora mostram a mensagem real
3.5.0 maio/2026: Linhas de resposta dissertativas injetadas inline no DOCX
                  da prova (após cada questão 9/10/11); linhas configuráveis
3.4.0 maio/2026: Download em PDF via Microsoft Word local; layout wide
3.3.1a maio/2026: Ajustes nos modelos Modelo_AF.docx e Modelo_AR.docx
3.3.1 maio/2026:  Suporte a gabarito Zipgrade; campos professor/disciplina
3.2.0 maio/2026:  Formatação markdown inline nas questões objetivas
3.0.0 maio/2026:  Migração para Streamlit; entrada via XLSX; download em ZIP
"""

import streamlit as st
import random
import os
import zipfile
import tempfile
import copy
import json
from io import BytesIO
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm
from datetime import datetime
import openpyxl
import re

# PDF — Windows/Mac via docx2pdf (Word); Linux via LibreOffice headless
import platform as _platform
import subprocess as _subprocess

_SISTEMA = _platform.system()   # "Windows", "Darwin" ou "Linux"

if _SISTEMA == "Linux":
    # Verifica se libreoffice está disponível no PATH
    try:
        _subprocess.run(["libreoffice", "--version"], capture_output=True, timeout=10)
        PDF_DISPONIVEL = True
    except Exception:
        PDF_DISPONIVEL = False
    _docx2pdf = None  # não usado no Linux
else:
    try:
        from docx2pdf import convert as _docx2pdf
        PDF_DISPONIVEL = True
    except ImportError:
        PDF_DISPONIVEL = False
        _docx2pdf = None

# Integração com Claude API — opcional
try:
    import anthropic
    ANTHROPIC_DISPONIVEL = True
except ImportError:
    ANTHROPIC_DISPONIVEL = False

# Integração com OpenAI DALL-E — opcional
try:
    from openai import OpenAI as _OpenAI
    import base64 as _base64
    OPENAI_DISPONIVEL = True
except ImportError:
    OPENAI_DISPONIVEL = False

# Reparador de JSON malformado (aspas/quebras de linha dentro de strings)
try:
    from json_repair import repair_json as _repair_json
    _JSONREPAIR = True
except ImportError:
    _JSONREPAIR = False

# --- Variações de instrução para evitar respostas determinísticas da IA ---

_VARIACOES_OBJ = [
    "Priorize questões que exijam raciocínio e análise, não apenas memorização.",
    "Dê preferência a questões que usem exemplos práticos ou situações-problema.",
    "Foque em questões que avaliem compreensão profunda de conceitos fundamentais.",
    "Priorize questões que envolvam comparação ou diferenciação de conceitos.",
    "Dê preferência a questões que testam aplicação de conhecimento em contextos novos.",
    "Selecione questões que envolvam interpretação de dados, gráficos ou cenários.",
    "Prefira questões que avaliem a capacidade de síntese e generalização do aluno.",
    "Foque em questões que abordem aspectos práticos e profissionais da disciplina.",
]

_VARIACOES_DIS = [
    "Crie questões que peçam ao aluno para EXPLICAR mecanismos, processos ou fenômenos.",
    "Crie questões que exijam COMPARAÇÃO entre conceitos, teorias ou abordagens distintas.",
    "Crie questões que peçam ANÁLISE CRÍTICA de situações-problema ou casos práticos.",
    "Crie questões que solicitem CONTEXTUALIZAÇÃO histórica, social ou profissional do tema.",
    "Crie questões que peçam ao aluno para RELACIONAR dois ou mais conceitos do conteúdo.",
    "Crie questões que desafiem o aluno a JUSTIFICAR uma afirmação, hipótese ou decisão.",
    "Crie questões que solicitem DESCRIÇÃO detalhada de um processo, etapa ou fenômeno.",
    "Crie questões que peçam ao aluno para AVALIAR vantagens e limitações de uma abordagem.",
    "Crie questões que exijam APLICAÇÃO prática de conceitos teóricos em cenários reais.",
    "Crie questões que envolvam SÍNTESE de múltiplos conceitos em uma resposta integrada.",
]

# --- Configurações Globais ---

SIMBOLOS_PROVA = {
    "A": "**",
    "B": "==",
    "C": "%%",
    "D": "//",
    "E": "++",
    "F": "##",
    "G": "!!",
    "H": "()"
}

LETRAS_PROVA = list(SIMBOLOS_PROVA.keys())
ABC = "ABCDE"

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MODELOS = {
    ("R", "Padrão"):   "Modelo_AR.docx",
    ("R", "Zipgrade"): "Modelo_ARZ.docx",
    ("F", "Padrão"):   "Modelo_AF.docx",
    ("F", "Zipgrade"): "Modelo_AFZ.docx",
}

_MD_PATTERN  = re.compile(r'\*\*\*(.*?)\*\*\*|\*\*(.*?)\*\*|\*(.*?)\*', re.DOTALL)
_IMG_TAG_PAT = re.compile(r'\[img:([^\]]+)\]', re.IGNORECASE)


def parse_img_tags(text):
    """Remove [img:nome] do texto e retorna (texto_limpo, nome_arquivo | None)."""
    match = _IMG_TAG_PAT.search(text)
    if not match:
        return text.strip(), None
    filename = match.group(1).strip()
    clean = _IMG_TAG_PAT.sub('', text).strip()
    return clean, filename


def parse_markdown_segments(text):
    """Divide text em segmentos com flags (bold, italic) baseadas em marcadores markdown."""
    segments = []
    last_end = 0
    for m in _MD_PATTERN.finditer(text):
        if m.start() > last_end:
            segments.append((text[last_end:m.start()], False, False))
        if m.group(1) is not None:
            segments.append((m.group(1), True, True))
        elif m.group(2) is not None:
            segments.append((m.group(2), True, False))
        else:
            segments.append((m.group(3), False, True))
        last_end = m.end()
    if last_end < len(text):
        segments.append((text[last_end:], False, False))
    if not segments:
        segments.append((text, False, False))
    return segments


# --- Funções de PDF ---

def docx_bytes_to_pdf_bytes(docx_bytes):
    """
    Converte bytes DOCX em bytes PDF.
    - Linux: LibreOffice headless via subprocess (Streamlit Cloud).
    - Windows/Mac: docx2pdf via Microsoft Word.
    Retorna (pdf_bytes, None) em caso de sucesso ou (None, msg_erro) em falha.
    """
    if not PDF_DISPONIVEL:
        return None, "Conversão de PDF não disponível neste ambiente."

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "entrada.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        if _SISTEMA == "Linux":
            # LibreOffice headless — disponível no Streamlit Cloud via packages.txt
            try:
                result = _subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "pdf",
                     "--outdir", tmpdir, docx_path],
                    capture_output=True, timeout=120
                )
                if result.returncode != 0:
                    return None, result.stderr.decode("utf-8", errors="replace") or "LibreOffice retornou erro."
                pdf_path = os.path.join(tmpdir, "entrada.pdf")
                with open(pdf_path, "rb") as f:
                    return f.read(), None
            except Exception as e:
                return None, str(e)
        else:
            # Windows / macOS — docx2pdf via Microsoft Word
            try:
                import pythoncom
                _com_disponivel = True
            except ImportError:
                _com_disponivel = False

            pdf_path = os.path.join(tmpdir, "saida.pdf")
            try:
                if _com_disponivel:
                    pythoncom.CoInitialize()
                _docx2pdf(docx_path, pdf_path)
                with open(pdf_path, "rb") as f:
                    return f.read(), None
            except Exception as e:
                return None, str(e)
            finally:
                if _com_disponivel:
                    try:
                        pythoncom.CoUninitialize()
                    except Exception:
                        pass


# --- Importação de documentos não estruturados via IA ---

def _json_para_texto(dados, nivel=0):
    """
    Converte recursivamente um dict/list JSON em texto legível para a IA.
    Trata especificamente o formato exportado pelo Google Forms API/Apps Script.
    """
    partes = []
    prefixo = "  " * nivel

    # Formato Google Forms API v1 (items com questionItem)
    if isinstance(dados, dict) and "items" in dados:
        titulo = dados.get("info", {}).get("title", "")
        if titulo:
            partes.append(f"Formulário: {titulo}")
        for item in dados["items"]:
            t = item.get("title", "").strip()
            if not t:
                continue
            qi = item.get("questionItem", {})
            q  = qi.get("question", {})
            cq = q.get("choiceQuestion", {})
            tq = q.get("textQuestion", {})

            if cq:
                partes.append(f"\nQuestão: {t}")
                opcoes = cq.get("options", [])
                corretas = [o["value"] for o in opcoes if o.get("isCorrect")]
                incorretas = [o["value"] for o in opcoes if not o.get("isCorrect")]
                if corretas:
                    partes.append(f"  Resposta correta: {corretas[0]}")
                for op in opcoes:
                    partes.append(f"  Alternativa: {op.get('value', '')}")
            elif tq or q.get("rowQuestion"):
                partes.append(f"\nQuestão dissertativa: {t}")
            else:
                partes.append(f"\n{t}")
        return "\n".join(partes)

    # Formato Apps Script serializado ou JSON genérico
    if isinstance(dados, dict):
        for k, v in dados.items():
            if isinstance(v, (dict, list)):
                sub = _json_para_texto(v, nivel + 1)
                if sub:
                    partes.append(f"{prefixo}{k}:\n{sub}")
            elif v is not None and str(v).strip():
                partes.append(f"{prefixo}{k}: {v}")
    elif isinstance(dados, list):
        for item in dados:
            sub = _json_para_texto(item, nivel)
            if sub:
                partes.append(sub)

    return "\n".join(partes)


def extrair_texto_arquivo(uploaded_file):
    """
    Extrai o conteúdo textual de um arquivo enviado pelo usuário.
    Suporta:
    - .docx  : parágrafos + tabelas
    - .xlsx  : todas as abas
    - .json  : Google Forms API export ou JSON genérico
    - .txt / .gs / outros : texto plano (UTF-8)
    Retorna (texto_str, None) em sucesso ou (None, msg_erro) em falha.
    """
    nome = uploaded_file.name.lower()
    conteudo = uploaded_file.read()
    uploaded_file.seek(0)

    try:
        if nome.endswith(".docx"):
            doc = Document(BytesIO(conteudo))
            partes = []
            for para in doc.paragraphs:
                txt = para.text.strip()
                if txt:
                    partes.append(txt)
            for table in doc.tables:
                for row in table.rows:
                    linha = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if linha:
                        partes.append(linha)
            return "\n".join(partes), None

        elif nome.endswith(".xlsx"):
            wb = openpyxl.load_workbook(BytesIO(conteudo), read_only=True, data_only=True)
            partes = []
            for ws in wb.worksheets:
                partes.append(f"[Planilha: {ws.title}]")
                for row in ws.iter_rows(values_only=True):
                    celulas = [str(c).strip() for c in row if c is not None and str(c).strip()]
                    if celulas:
                        partes.append(" | ".join(celulas))
            wb.close()
            return "\n".join(partes), None

        elif nome.endswith(".json"):
            texto_bruto = conteudo.decode("utf-8", errors="replace")
            try:
                dados = json.loads(texto_bruto)
                return _json_para_texto(dados), None
            except json.JSONDecodeError:
                # Se não parsear, manda como texto mesmo
                return texto_bruto, None

        else:  # .txt, .gs, .csv ou qualquer texto plano
            texto = conteudo.decode("utf-8", errors="replace")
            return texto, None

    except Exception as e:
        return None, str(e)


def _parse_json_resposta(texto):
    """
    Extrai e parseia o bloco JSON da resposta do modelo.
    Tenta 3 estratégias em sequência:
    1. json.loads direto (resposta já é JSON puro)
    2. Extração via regex + json.loads
    3. json-repair (corrige aspas/vírgulas/quebras de linha malformadas)
    Retorna (dict, None) ou (None, msg_erro).
    """
    # Remove blocos de código markdown se presentes
    limpo = re.sub(r'```(?:json)?\s*', '', texto).strip()
    limpo = re.sub(r'```\s*$', '', limpo).strip()

    # Tentativa 1: texto inteiro já é JSON válido
    try:
        return json.loads(limpo), None
    except json.JSONDecodeError:
        pass

    # Tentativa 2: extrai o maior bloco {…} e tenta parsear
    match = re.search(r'\{[\s\S]*\}', limpo)
    if match:
        bloco = match.group(0)
        try:
            return json.loads(bloco), None
        except json.JSONDecodeError:
            # Tentativa 3: json-repair conserta aspas/vírgulas/newlines dentro de strings
            if _JSONREPAIR:
                try:
                    reparado = _repair_json(bloco, return_objects=True)
                    if isinstance(reparado, dict):
                        return reparado, None
                except Exception:
                    pass
            # Sem json-repair: retorna erro descritivo
            return None, (
                "A resposta da API contém JSON malformado (provavelmente aspas dentro "
                "do texto das questões). Tente novamente — o modelo às vezes corrige sozinho."
            )

    return None, "A API não retornou um bloco JSON reconhecível. Tente novamente."


def _chamar_api_claude(prompt, api_key, max_tokens=8192):
    """
    Chamada base à API Claude. Retorna (dict, None) ou (None, erro).
    Centraliza autenticação, chamada e parsing robusto do JSON.
    """
    try:
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}]
        )
        resposta = message.content[0].text.strip()
        return _parse_json_resposta(resposta)
    except Exception as e:
        return None, f"Erro na chamada à API: {e}"


def _gerar_imagem_dalle(prompt, api_key_openai):
    """
    Gera uma imagem via gpt-image-1 e retorna (bytes_png, None) ou (None, erro).
    Compatível com SDK openai >= 2.x (gpt-image-1) e >= 1.x (dall-e-3).
    """
    try:
        import urllib.request as _urlreq
        client = _OpenAI(api_key=api_key_openai)
        response = client.images.generate(
            model="gpt-image-1",
            prompt=prompt,
            size="1024x1024",
            quality="auto",
            n=1,
        )
        data = response.data[0]
        # gpt-image-1 retorna base64 diretamente em data.b64
        if getattr(data, "b64", None):
            return _base64.b64decode(data.b64), None
        # dall-e-3 / fallback: b64_json
        if getattr(data, "b64_json", None):
            return _base64.b64decode(data.b64_json), None
        # Fallback URL
        if getattr(data, "url", None):
            with _urlreq.urlopen(data.url) as r:
                return r.read(), None
        return None, "Nenhum dado de imagem na resposta da API."
    except Exception as e:
        return None, str(e)


def _gerar_prompts_imagem_claude(questoes_obj, questoes_dis, api_key_claude):
    """
    Pede ao Claude um prompt DALL-E para cada questão aprovada.
    Retorna (dict {enunciado: prompt_dalle}, None) ou (None, erro).
    Claude decide se cada questão se beneficia de imagem; retorna null para as que não precisam.
    """
    lista_q = []
    for i, q in enumerate(questoes_obj or [], start=1):
        lista_q.append(f"OBJ {i}: {q[0][:300]}")
    for i, q in enumerate(questoes_dis or [], start=1):
        lista_q.append(f"DIS {i}: {q[:300]}")

    if not lista_q:
        return {}, None

    questoes_txt = "\n".join(lista_q)

    prompt = f"""Você é um especialista em criação de recursos visuais para provas universitárias.

Para CADA questão abaixo, crie um prompt em INGLÊS para gerar uma imagem ilustrativa com DALL-E.
A imagem deve representar visualmente o tema central da questão — pode ser um diagrama, infográfico, cena, objeto, processo ou conceito.
TODAS as questões devem receber um prompt. Não use null.

QUESTÕES:
{questoes_txt}

Retorne EXCLUSIVAMENTE um JSON válido com exatamente {len(lista_q)} itens:
{{
  "imagens": [
    {{"enunciado_prefix": "OBJ 1", "prompt_dalle": "..." }},
    {{"enunciado_prefix": "OBJ 2", "prompt_dalle": "..." }},
    {{"enunciado_prefix": "DIS 1", "prompt_dalle": "..." }}
  ]
}}

Regras para cada prompt DALL-E:
- Escreva em inglês
- Descreva uma cena ou diagrama visual claro e relevante para o tema da questão
- Estilo sugerido: "educational illustration, clean design, white background, no text"
- NÃO inclua texto, letras ou legendas na imagem gerada"""

    resultado, erro = _chamar_api_claude(prompt, api_key_claude, max_tokens=2048)
    if erro:
        return None, erro

    mapa = {}
    todas_q = list(questoes_obj or []) + list(questoes_dis or [])
    prefixos_obj = [f"OBJ {i}" for i in range(1, len(questoes_obj or []) + 1)]
    prefixos_dis = [f"DIS {i}" for i in range(1, len(questoes_dis or []) + 1)]
    prefixos = prefixos_obj + prefixos_dis

    for item in resultado.get("imagens", []):
        pref  = item.get("enunciado_prefix", "")
        dalle = item.get("prompt_dalle")
        if dalle and pref in prefixos:
            idx = prefixos.index(pref)
            if idx < len(todas_q):
                enunciado = todas_q[idx][0] if isinstance(todas_q[idx], tuple) else todas_q[idx]
                mapa[enunciado] = dalle

    return mapa, None


def _prompt_objetivas(texto, ja_aprovadas=None, n_faltando=None):
    """
    Prompt dedicado à EXTRAÇÃO de questões objetivas (múltipla escolha).
    Usa apenas o conteúdo que já está no documento.
    Injeta uma instrução aleatória para variar a seleção entre chamadas.
    """
    ctx = ""
    if ja_aprovadas:
        lista = "\n".join(f"  - {q[0][:120]}" for q in ja_aprovadas)
        ctx = f"\nQuestões JÁ APROVADAS — NÃO repita:\n{lista}\n"

    qtd = ""
    if n_faltando:
        qtd = (
            f"\nPreciso de EXATAMENTE {n_faltando} questão(ões) objetiva(s) novas "
            "(não repetidas). Se não encontrar no documento, crie questões plausíveis sobre o mesmo tema.\n"
        )

    variacao = random.choice(_VARIACOES_OBJ)

    return f"""Você é um assistente especializado em provas universitárias brasileiras.
Extraia as questões de múltipla escolha presentes no documento abaixo.
Dica de foco para esta extração: {variacao}
{ctx}{qtd}
REGRAS OBRIGATÓRIAS:
1. Extraia no máximo 10 questões objetivas — priorize as mais representativas do conteúdo.
2. Para cada questão inclua: enunciado completo, resposta correta e EXATAMENTE 4 distratores.
3. Se houver gabarito indicado (ex: 'Gabarito: B', 'Resposta: C'), use-o para identificar a correta.
4. Retorne o TEXTO de cada alternativa, sem a letra (A, B, C...).
5. Se houver menos de 4 distratores no documento, crie um adicional plausível sobre o mesmo tema.
6. Ignore: cabeçalhos, nome da instituição, curso, data, tempo de prova, instruções gerais.
7. Aspas duplas (") dentro de textos devem ser substituídas por aspas simples (').

Retorne EXCLUSIVAMENTE um JSON válido:
{{
  "objetivas": [
    {{
      "enunciado": "texto completo do enunciado",
      "resposta_correta": "texto da alternativa correta (sem letra)",
      "distratores": ["distrator 1", "distrator 2", "distrator 3", "distrator 4"]
    }}
  ],
  "dissertativas": []
}}

DOCUMENTO:
{texto[:22000]}"""


def _prompt_dissertativas(texto, temas_obj=None, ja_aprovadas=None, n_faltando=5):
    """
    Prompt dedicado à CRIAÇÃO de questões dissertativas com base nos temas do documento.

    Documentos de Google Forms tipicamente têm apenas múltipla escolha.
    Por isso este prompt instrui o modelo a CRIAR questões abertas adequadas
    para o nível universitário, baseando-se nos temas e conceitos identificados.
    Gera um pool de até 5 questões (padrão) para seleção aleatória entre versões.
    Injeta uma instrução aleatória para variar a abordagem entre chamadas.
    """
    ctx_temas = ""
    if temas_obj:
        def _enunciado(q):
            # Aceita tanto dict (API bruta) quanto tupla (já convertida)
            return q["enunciado"][:100] if isinstance(q, dict) else q[0][:100]
        resumo = "\n".join(f"  - {_enunciado(q)}" for q in temas_obj[:10])
        ctx_temas = f"\nTemas identificados pelas questões objetivas já extraídas:\n{resumo}\n"

    ctx_aprovadas = ""
    if ja_aprovadas:
        lista = "\n".join(f"  - {q[:120]}" for q in ja_aprovadas)
        ctx_aprovadas = f"\nQuestões dissertativas JÁ APROVADAS — NÃO repita:\n{lista}\n"

    variacao = random.choice(_VARIACOES_DIS)

    return f"""Você é um professor universitário elaborando questões dissertativas para uma prova.

O documento abaixo é um banco de questões objetivas de uma disciplina universitária.
Seu papel é criar questões DISSERTATIVAS (abertas) adequadas para a mesma disciplina e nível.
Abordagem para esta rodada: {variacao}
{ctx_temas}{ctx_aprovadas}
IMPORTANTE: o documento provavelmente NÃO contém questões dissertativas explícitas —
você deve CRIAR questões originais baseando-se nos temas e conceitos presentes.

REGRAS para criar as questões dissertativas:
1. Crie EXATAMENTE {n_faltando} questões dissertativas originais e DIVERSAS entre si.
2. Use verbos de comando adequados ao nível universitário:
   Explique, Descreva, Analise, Justifique, Discorra sobre, Compare, Relacione,
   Elabore, Caracterize, Discuta, Avalie, Demonstre, Apresente, Contextualize.
3. Cada questão deve ser autocontida (o aluno consegue responder sem ler as outras).
4. Nível: exigem raciocínio crítico, síntese e elaboração textual — não respostas sim/não.
5. Baseie-se nos temas e conceitos do documento — não invente assuntos fora do escopo.
6. As questões devem cobrir DIFERENTES aspectos/temas do documento — não repita o mesmo tema.
7. Formato final: apenas o enunciado da questão, sem gabarito, critérios ou pontuação.
   (Este enunciado será inserido diretamente na prova com espaço para resposta manuscrita.)
8. Aspas duplas (") dentro de textos devem ser substituídas por aspas simples (').

Retorne EXCLUSIVAMENTE um JSON válido:
{{
  "objetivas": [],
  "dissertativas": [
    {{
      "enunciado": "Enunciado completo da questão dissertativa criada.",
      "resposta_esperada": "Resposta modelo resumida (3 a 6 linhas) que o professor usará como gabarito."
    }}
  ]
}}

DOCUMENTO (banco de questões da disciplina):
{texto[:22000]}"""


def processar_com_claude(texto, api_key, tipo_questoes="ambas"):
    """
    Processa o documento em até duas chamadas à API:

    - "objetivas"     → extrai todas as questões de múltipla escolha.
    - "dissertativas" → cria questões abertas com base nos temas do documento.
    - "ambas"         → faz as duas chamadas separadas e combina o resultado.

    A separação garante que cada chamada usa o prompt mais adequado para cada tipo,
    evitando que o modelo "esqueça" as dissertativas ao priorizar as objetivas.
    Retorna (dict, None) ou (None, msg_erro).
    """
    resultado = {"objetivas": [], "dissertativas": []}

    if tipo_questoes in ("objetivas", "ambas"):
        prompt_obj = _prompt_objetivas(texto)
        res_obj, erro = _chamar_api_claude(prompt_obj, api_key)
        if erro:
            return None, erro
        resultado["objetivas"] = res_obj.get("objetivas", [])

    if tipo_questoes in ("dissertativas", "ambas"):
        temas = resultado.get("objetivas") or []
        prompt_dis = _prompt_dissertativas(texto, temas_obj=temas)  # usa n_faltando=5
        res_dis, erro = _chamar_api_claude(prompt_dis, api_key)
        if erro:
            return None, erro
        resultado["dissertativas"] = res_dis.get("dissertativas", [])

    return resultado, None


def regenerar_nao_confirmadas(texto, api_key, tipo_questoes,
                               obj_confirmadas, dis_confirmadas,
                               n_obj_faltando, n_dis_faltando):
    """
    Regenera apenas as questões não confirmadas usando os prompts dedicados.
    Passa as já confirmadas como contexto para evitar repetição.
    Retorna (dict, None) ou (None, msg_erro).
    """
    resultado = {"objetivas": [], "dissertativas": []}

    if n_obj_faltando > 0:
        prompt = _prompt_objetivas(texto, ja_aprovadas=obj_confirmadas,
                                   n_faltando=n_obj_faltando)
        res, erro = _chamar_api_claude(prompt, api_key)
        if erro:
            return None, erro
        resultado["objetivas"] = res.get("objetivas", [])

    if n_dis_faltando > 0:
        prompt = _prompt_dissertativas(texto, temas_obj=obj_confirmadas,
                                       ja_aprovadas=dis_confirmadas,
                                       n_faltando=n_dis_faltando)
        res, erro = _chamar_api_claude(prompt, api_key)
        if erro:
            return None, erro
        resultado["dissertativas"] = res.get("dissertativas", [])

    return resultado, None


def _converter_resultado_ia(resultado):
    """
    Converte o dict retornado pela API para as listas internas do gerador.
    Retorna (lista_objetivas, lista_dissertativas, lista_gabaritos_dis).
    - lista_objetivas      : list of tuples (enunciado, resp_correta, alt1..alt5)
    - lista_dissertativas  : list of strings (enunciados)
    - lista_gabaritos_dis  : list of strings (respostas esperadas, mesma ordem)
    """
    objetivas = []
    for q in resultado.get("objetivas", []):
        try:
            dist = list(q.get("distratores", []))[:4]
            while len(dist) < 4:
                dist.append("(distrator não identificado)")
            alternativas = [q["resposta_correta"]] + dist
            objetivas.append((q["enunciado"], q["resposta_correta"], *alternativas))
        except (KeyError, TypeError):
            continue

    dissertativas = []
    gabaritos_dis = []
    for q in resultado.get("dissertativas", []):
        if isinstance(q, dict) and q.get("enunciado"):
            dissertativas.append(q["enunciado"])
            gabaritos_dis.append(q.get("resposta_esperada", ""))

    return objetivas, dissertativas, gabaritos_dis


# --- Manipulação de seções e linhas de resposta (XML direto) ---

def _separar_secao_dissertativas(document):
    """
    Separa as questões dissertativas (9/10/11) para uma seção própria de
    1 coluna (largura total da página).

    O template AR usa 2 colunas para toda a seção de questões (1–11).
    Esta função:
    1. Copia o sectPr de 2 colunas para o parágrafo vazio antes da Q9,
       usando tipo 'continuous' para que o Word BALANCEIE automaticamente
       as colunas das objetivas (sem deixar a coluna direita vazia).
    2. Adiciona 'pageBreakBefore' ao parágrafo da Q9 para que as
       dissertativas sempre iniciem em página nova.
    3. Altera o sectPr original para 1 coluna (sem tipo forçado),
       criando a seção de largura total para as dissertativas.

    Se o template não tiver layout de 2 colunas, retorna sem modificar nada.
    """
    paragrafos = list(document.paragraphs)

    # 1. Localiza o sectPr que define a seção de 2 colunas
    sectPr_2col = None
    for p in paragrafos:
        sp = p._p.find(f".//{qn('w:sectPr')}")
        if sp is not None:
            cols = sp.find(qn("w:cols"))
            if cols is not None and cols.get(qn("w:num"), "1") == "2":
                sectPr_2col = sp
                break

    if sectPr_2col is None:
        return  # Sem layout 2-col; nada a fazer

    # 2. Localiza a primeira questão dissertativa (texto começa com "9. (")
    idx_q9 = None
    for i, p in enumerate(paragrafos):
        if p.text.lstrip().startswith("9. ("):
            idx_q9 = i
            break

    if idx_q9 is None or idx_q9 == 0:
        return  # Q9 não encontrada

    # 3. Parágrafo separador (vazio antes da Q9) e a própria Q9
    para_sep = paragrafos[idx_q9 - 1]
    para_q9  = paragrafos[idx_q9]

    # 4. Copia o sectPr 2-col para o parágrafo separador com tipo 'continuous'.
    #    'continuous' faz o Word BALANCEAR as colunas das objetivas antes de
    #    encerrar a seção — evita coluna direita vazia ou linhas em branco.
    novo_sectPr = copy.deepcopy(sectPr_2col)
    tipo_el = novo_sectPr.find(qn("w:type"))
    if tipo_el is None:
        tipo_el = OxmlElement("w:type")
        novo_sectPr.insert(0, tipo_el)
    tipo_el.set(qn("w:val"), "continuous")

    pPr_sep = para_sep._p.get_or_add_pPr()
    for old in pPr_sep.findall(qn("w:sectPr")):
        pPr_sep.remove(old)
    pPr_sep.append(novo_sectPr)

    # 5. Converte o sectPr original para 1 coluna (seção das dissertativas)
    cols_orig = sectPr_2col.find(qn("w:cols"))
    if cols_orig is not None:
        sectPr_2col.remove(cols_orig)
    cols_1 = OxmlElement("w:cols")
    cols_1.set(qn("w:num"), "1")
    sectPr_2col.append(cols_1)

    # Remove qualquer tipo herdado do sectPr original
    tipo_orig = sectPr_2col.find(qn("w:type"))
    if tipo_orig is not None:
        sectPr_2col.remove(tipo_orig)

    # 6. Adiciona pageBreakBefore à Q9 → dissertativas sempre em página nova
    pPr_q9 = para_q9._p.get_or_add_pPr()
    pb = OxmlElement("w:pageBreakBefore")
    pPr_q9.insert(0, pb)


def _criar_p_linha_resposta(font_size_pt=11):
    """
    Cria um elemento XML <w:p> com borda inferior estilo caderno.
    Não altera parágrafos existentes — cria um elemento novo para inserção.
    """
    new_p = OxmlElement("w:p")

    pPr = OxmlElement("w:pPr")

    # Espaçamento: sem espaço antes, 4 pt depois
    pSpacing = OxmlElement("w:spacing")
    pSpacing.set(qn("w:before"), "0")
    pSpacing.set(qn("w:after"), "80")   # 4pt × 20 twips
    pPr.append(pSpacing)

    # Borda inferior (linha de caderno)
    pBdr = OxmlElement("w:pBdr")
    borda = OxmlElement("w:bottom")
    borda.set(qn("w:val"),   "single")
    borda.set(qn("w:sz"),    "6")
    borda.set(qn("w:space"), "1")
    borda.set(qn("w:color"), "808080")
    pBdr.append(borda)
    pPr.append(pBdr)

    new_p.append(pPr)

    # Run com NBSP para garantir altura de linha
    r_el = OxmlElement("w:r")
    rPr  = OxmlElement("w:rPr")
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), str(int(font_size_pt * 2)))   # half-points
        rPr.append(el)
    r_el.append(rPr)

    t_el = OxmlElement("w:t")
    t_el.text = "\u00a0"   # non-breaking space
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_el.append(t_el)
    new_p.append(r_el)

    return new_p


def _injetar_linhas_apos(para, n_linhas, font_size_pt=11):
    """
    Insere n_linhas de resposta logo após o parágrafo 'para'.
    keepNext é adicionado à questão para não separá-la das linhas.
    O parágrafo vazio do template que segue as linhas (separador entre
    questões) tem seu espaçamento zerado para evitar espaço em branco extra.
    """
    pPr = para._p.get_or_add_pPr()
    kwn = OxmlElement("w:keepNext")
    pPr.insert(0, kwn)

    ultimo = para._p
    for _ in range(n_linhas):
        nova = _criar_p_linha_resposta(font_size_pt)
        ultimo.addnext(nova)
        ultimo = nova

    # Zera o parágrafo vazio do template que fica após as linhas
    # (separador Q9→Q10, Q10→Q11 no template AR) — evita espaço em branco extra
    proximo = ultimo.getnext()
    if proximo is not None and proximo.tag == ultimo.tag:
        pPr_prox = proximo.find(qn("w:pPr"))
        if pPr_prox is None:
            pPr_prox = OxmlElement("w:pPr")
            proximo.insert(0, pPr_prox)
        sp = pPr_prox.find(qn("w:spacing"))
        if sp is None:
            sp = OxmlElement("w:spacing")
            pPr_prox.append(sp)
        sp.set(qn("w:before"), "0")
        sp.set(qn("w:after"), "0")
        sp.set(qn("w:line"), "20")        # 1 twip = quase invisível
        sp.set(qn("w:lineRule"), "exact")


def _injetar_imagem_apos(document, para, img_bytes, largura_cm=6.0):
    """
    Insere um parágrafo centralizado com imagem imediatamente após 'para'.
    Retorna o novo parágrafo (wrapper python-docx).
    """
    new_p_el = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    new_p_el.append(pPr)
    para._p.addnext(new_p_el)

    # Localiza o wrapper Python para o elemento recém-inserido
    new_para = None
    for p in document.paragraphs:
        if p._p is new_p_el:
            new_para = p
            break
    if new_para is None:
        return para

    run = new_para.add_run()
    run.add_picture(BytesIO(img_bytes), width=Cm(largura_cm))
    return new_para


# --- Funções de Lógica Central ---

def replace_text_in_paragraph_runs(paragraph, old_text, new_text, bold_prefix=False):
    """
    Substitui um texto em um parágrafo, lidando com múltiplas 'runs'.
    Preserva a formatação ao redor e aplica negrito ao prefixo (A), (B)...
    se bold_prefix for True.
    """
    if old_text not in paragraph.text:
        return False

    if not bold_prefix:
        replaced = False
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                replaced = True
        return replaced

    char_fmt = []
    for run in paragraph.runs:
        fmt = (run.bold, run.italic, run.underline, run.font.name, run.font.size)
        char_fmt.extend([fmt] * len(run.text))

    ref_run = None
    for run in paragraph.runs:
        if ref_run is None:
            ref_run = run
        if old_text in run.text:
            ref_run = run
            break

    original_text   = "".join(run.text for run in paragraph.runs)
    placeholder_pos = original_text.find(old_text)
    before_text     = original_text[:placeholder_pos]
    after_text      = original_text[placeholder_pos + len(old_text):]

    paragraph.clear()

    def add_run(text):
        r = paragraph.add_run(text)
        if ref_run is not None:
            r.bold    = ref_run.bold
            r.italic  = ref_run.italic
            r.underline = ref_run.underline
            if ref_run.font.name:
                r.font.name = ref_run.font.name
            if ref_run.font.size:
                r.font.size = ref_run.font.size
        return r

    def add_run_with_markdown(text):
        for seg_text, md_bold, md_italic in parse_markdown_segments(text):
            if not seg_text:
                continue
            r = add_run(seg_text)
            r.bold   = r.bold   or md_bold
            r.italic = r.italic or md_italic

    def add_formatted_segment(text, start_pos):
        if not text:
            return
        i = 0
        while i < len(text):
            fmt = char_fmt[start_pos + i]
            j = i + 1
            while j < len(text) and char_fmt[start_pos + j] == fmt:
                j += 1
            bold, italic, underline, font_name, font_size = fmt
            r = paragraph.add_run(text[i:j])
            r.bold      = bold
            r.italic    = italic
            r.underline = underline
            if font_name:
                r.font.name = font_name
            if font_size:
                r.font.size = font_size
            i = j

    add_formatted_segment(before_text, 0)

    temp = new_text
    while True:
        match = False
        for char in ABC:
            prefix = f"({char})"
            if prefix in temp:
                before, temp = temp.split(prefix, 1)
                if before:
                    add_run_with_markdown(before)
                bold_run = add_run(prefix)
                bold_run.bold = True
                match = True
                break
        if not match:
            if temp:
                add_run_with_markdown(temp)
            break

    add_formatted_segment(after_text, placeholder_pos + len(old_text))
    return True


def all_paragraphs(document):
    """
    Itera sobre todos os parágrafos do documento: corpo, tabelas, cabeçalhos,
    rodapés e Structured Document Tags (w:sdt), percorrendo recursivamente.
    """
    _CONTAINER_TAGS = {
        qn("w:body"), qn("w:tbl"), qn("w:tr"), qn("w:tc"),
        qn("w:hdr"), qn("w:ftr"),
        qn("w:sdt"), qn("w:sdtContent"),
    }

    def _iter(element):
        for child in element:
            if child.tag == qn("w:p"):
                from docx.text.paragraph import Paragraph
                yield Paragraph(child, element)
            elif child.tag in _CONTAINER_TAGS:
                yield from _iter(child)

    yield from _iter(document.element.body)
    for section in document.sections:
        yield from _iter(section.header._element)
        yield from _iter(section.footer._element)


def criar_prova(nome_prova, simbolo_rodape, qt_questoes, questoes_selecionadas):
    """
    Cria o conteúdo de uma prova (questões objetivas formatadas e gabarito).
    Retorna (lista_de_textos_de_questões, string_gabarito, lista_enunciados).
    """
    lista_questoes = []
    lista_enunciados = []
    gabarito = "GABARITO\n"

    for indice, questao_tupla in enumerate(questoes_selecionadas, start=1):
        pergunta_texto         = questao_tupla[0]
        resposta_correta_texto = questao_tupla[1]
        opcoes_originais       = list(questao_tupla[2:])

        questao_texto_completo = pergunta_texto + "\n\n"

        opcoes_com_indice = [(opcoes_originais[j], j) for j in range(len(opcoes_originais))]
        random.shuffle(opcoes_com_indice)

        posicao_resposta_correta = -1
        for idx_exib, (opcao_texto, _) in enumerate(opcoes_com_indice):
            questao_texto_completo += f"({ABC[idx_exib]}) {opcao_texto}\n"
            if opcao_texto.lower().replace(" ", "") == resposta_correta_texto.lower().replace(" ", ""):
                posicao_resposta_correta = idx_exib

        lista_questoes.append(questao_texto_completo)
        lista_enunciados.append(pergunta_texto)

        if posicao_resposta_correta != -1:
            gabarito += f"{indice}: {ABC[posicao_resposta_correta]}\n"
        else:
            gabarito += f"{indice}: ERRO (Resposta não identificada para '{pergunta_texto[:40]}...')\n"

    return lista_questoes, gabarito, lista_enunciados


# --- Funções de I/O ---

def get_questoes_xlsx(uploaded_file):
    """
    Lê questões objetivas de um arquivo XLSX.
    Espera: A=Enunciado, B=Resposta correta, C–F=Distratores, G=Imagem (opcional).
    Retorna (questoes, imagens_dict) onde imagens_dict = {enunciado: nome_arquivo}.
    """
    questoes = []
    imagens = {}
    try:
        wb = openpyxl.load_workbook(uploaded_file, read_only=True, data_only=True)
        ws = wb.active
        for i, row in enumerate(ws.iter_rows(values_only=True), start=1):
            if not row or all(cell is None for cell in row):
                continue
            if len(row) < 6:
                st.warning(f"Linha {i}: menos de 6 colunas. Ignorando.")
                continue
            pergunta_raw     = str(row[0]).strip() if row[0] is not None else ""
            pergunta, img_inline = parse_img_tags(pergunta_raw)
            resposta_correta = str(row[1]).strip() if row[1] is not None else ""
            distratores = [
                str(row[j]).strip() if row[j] is not None else ""
                for j in range(2, 6)
            ]
            if not all([pergunta, resposta_correta] + distratores):
                st.warning(f"Linha {i}: célula vazia em coluna essencial (A–F). Ignorando.")
                continue
            alternativas = [resposta_correta] + distratores
            questoes.append((pergunta, resposta_correta, *alternativas))
            # Coluna G (índice 6) tem prioridade; [img:] inline é fallback
            if len(row) > 6 and row[6] is not None:
                nome_img = str(row[6]).strip()
                if nome_img:
                    imagens[pergunta] = nome_img
            elif img_inline:
                imagens[pergunta] = img_inline
        wb.close()
    except Exception as e:
        st.error(f"Erro ao ler questões objetivas: {e}")
        return None, {}
    return (questoes or None), imagens


def get_questoes_dissertativas_xlsx(uploaded_file):
    """
    Lê questões discursivas de um arquivo XLSX.
    Espera: A=Enunciado, B=Imagem (opcional).
    Retorna (questoes, imagens_dict) onde imagens_dict = {enunciado: nome_arquivo}.
    """
    questoes = []
    imagens = {}
    try:
        wb = openpyxl.load_workbook(uploaded_file, read_only=True, data_only=True)
        ws = wb.active
        for i, row in enumerate(ws.iter_rows(values_only=True), start=1):
            if not row or row[0] is None:
                continue
            enunciado_raw = str(row[0]).strip()
            enunciado, img_inline = parse_img_tags(enunciado_raw)
            if enunciado:
                questoes.append(enunciado)
                # Coluna B (índice 1) tem prioridade; [img:] inline é fallback
                if len(row) > 1 and row[1] is not None:
                    nome_img = str(row[1]).strip()
                    if nome_img:
                        imagens[enunciado] = nome_img
                elif img_inline:
                    imagens[enunciado] = img_inline
            else:
                st.warning(f"Linha {i}: enunciado vazio. Ignorando.")
        wb.close()
    except Exception as e:
        st.error(f"Erro ao ler questões discursivas: {e}")
        return None, {}
    return (questoes or None), imagens


def _injetar_pontuacao_cabecalho(document, pontos_obj):
    """
    Localiza o parágrafo 'QUESTÕES' no template e substitui por
    'Questões Objetivas = X pts' (se pontos_obj preenchido) ou
    'Questões Objetivas' (sempre — funciona com e sem pontuação).
    """
    novo_texto = f"Questões Objetivas = {pontos_obj}" if pontos_obj else "Questões Objetivas"
    for para in document.paragraphs:
        txt = para.text.strip()
        if txt.upper() in ("QUESTÕES", "QUESTOES"):
            replace_text_in_paragraph_runs(para, txt, novo_texto)
            break


def _injetar_titulo_dissertativas(document, para_primeira_dis, pontos_dis):
    """
    Localiza o parágrafo vazio imediatamente antes da primeira questão
    dissertativa (o separador de seção) e replica nele o estilo completo
    do heading de objetivas (cinza D9D9D9, centralizado, negrito, NormalWeb),
    preservando o sectPr já inserido por _separar_secao_dissertativas.
    Compara por _p (elemento XML) pois document.paragraphs cria novos
    wrappers Python a cada acesso.
    """
    paragrafos = list(document.paragraphs)
    idx = None
    for i, p in enumerate(paragrafos):
        if p._p is para_primeira_dis._p:
            idx = i
            break
    if idx is None or idx == 0:
        return

    para_sep = paragrafos[idx - 1]
    texto = f"Questões Discursivas = {pontos_dis}" if pontos_dis else "Questões Discursivas"

    # Salva o sectPr que _separar_secao_dissertativas adicionou ao separador
    pPr_sep = para_sep._p.find(qn("w:pPr"))
    sectPr_salvo = None
    if pPr_sep is not None:
        s = pPr_sep.find(qn("w:sectPr"))
        if s is not None:
            sectPr_salvo = copy.deepcopy(s)

    # Localiza o parágrafo heading das objetivas para copiar o pPr completo
    para_heading_obj = None
    for p in paragrafos:
        if "Objetivas" in p.text or p.text.strip().upper() in ("QUESTÕES", "QUESTOES"):
            para_heading_obj = p
            break

    if para_heading_obj is not None:
        pPr_obj = para_heading_obj._p.find(qn("w:pPr"))
        if pPr_obj is not None:
            # Copia o pPr do heading e reinsere o sectPr do separador
            novo_pPr = copy.deepcopy(pPr_obj)
            for s in novo_pPr.findall(qn("w:sectPr")):
                novo_pPr.remove(s)
            if sectPr_salvo is not None:
                novo_pPr.append(sectPr_salvo)
            # Substitui o pPr do separador
            if pPr_sep is not None:
                para_sep._p.remove(pPr_sep)
            para_sep._p.insert(0, novo_pPr)

        # Remove runs atuais e copia o run do heading (com formatação idêntica)
        for r in para_sep._p.findall(qn("w:r")):
            para_sep._p.remove(r)
        for r_orig in para_heading_obj._p.findall(qn("w:r")):
            novo_r = copy.deepcopy(r_orig)
            for t in novo_r.findall(qn("w:t")):
                t.text = texto
            para_sep._p.append(novo_r)
            break  # apenas o primeiro run
    else:
        # Fallback: sem heading de referência, aplica bold simples
        if para_sep.runs:
            run = para_sep.runs[0]
            run.text = texto
            run.bold = True
        else:
            para_sep.add_run(texto).bold = True


def gera_prova_bytes(modelo_caminho, identificador_prova, questoes_formatadas,
                     simbolo_rodape, tipo_avaliacao, questoes_dissertativas=None,
                     professor="", disciplina="", linhas_por_questao=8,
                     pontos_obj="", pontos_dis="",
                     imagens_questoes=None, enunciados_obj=None):
    """
    Preenche o template DOCX com as questões e retorna os bytes do arquivo gerado.

    Para provas Regimentais (AR):
    - As questões dissertativas são inseridas em uma seção de 1 coluna
      (largura total da página), separada da seção de 2 colunas das objetivas,
      sempre iniciando em página nova.
    - Linhas de resposta são injetadas inline após cada questão dissertativa.
    - Se pontos_obj/pontos_dis forem fornecidos, o cabeçalho 'QUESTÕES' é
      substituído pelo texto com pontuação.
    """
    try:
        document = Document(modelo_caminho)

        # Injeta cabeçalho "Questões Objetivas [= X pts]" — sempre
        _injetar_pontuacao_cabecalho(document, pontos_obj)

        # Substitui placeholders das questões objetivas
        _enunciados = enunciados_obj or []
        for i, questao_texto in enumerate(questoes_formatadas, start=1):
            placeholder = f"{{{{Questão {i} aqui}}}}"
            para_encontrado = None
            for paragraph in document.paragraphs:
                if replace_text_in_paragraph_runs(paragraph, placeholder, questao_texto, bold_prefix=True):
                    para_encontrado = paragraph
                    break
            if para_encontrado and imagens_questoes and i <= len(_enunciados):
                enunciado_i = _enunciados[i - 1]
                if enunciado_i in imagens_questoes:
                    _injetar_imagem_apos(document, para_encontrado,
                                        imagens_questoes[enunciado_i], largura_cm=5.5)

        # Substitui questões discursivas e guarda referência a cada parágrafo
        _N_SLOTS_DIS = 3   # número de slots no template AR (questões 9, 10, 11)
        paragrafos_dissertativos = []   # parágrafos do texto da questão (para título/seção)
        paragrafos_ultimos_dis   = []   # último parágrafo por questão (questão ou imagem)
        if questoes_dissertativas:
            # Atualiza o valor por questão nos slots do template ANTES de preencher
            n_dis = len(questoes_dissertativas)
            val_dis = 3.0 / n_dis
            novo_val_pt = f"({f'{val_dis:.2f}'.replace('.', ',')} pt.)"
            for idx in range(9, 9 + _N_SLOTS_DIS):
                ph = f"{{{{Questão {idx} aqui}}}}"
                for paragraph in document.paragraphs:
                    if ph in paragraph.text and "(1,00 pt.)" in paragraph.text:
                        replace_text_in_paragraph_runs(paragraph, "(1,00 pt.)", novo_val_pt)
                        break

            for idx, texto_diss in enumerate(questoes_dissertativas, start=9):
                placeholder_diss = f"{{{{Questão {idx} aqui}}}}"
                for paragraph in document.paragraphs:
                    if replace_text_in_paragraph_runs(paragraph, placeholder_diss, texto_diss, bold_prefix=True):
                        paragrafos_dissertativos.append(paragraph)
                        last_para = paragraph
                        if imagens_questoes and texto_diss in imagens_questoes:
                            last_para = _injetar_imagem_apos(
                                document, paragraph,
                                imagens_questoes[texto_diss], largura_cm=5.5
                            )
                        paragrafos_ultimos_dis.append(last_para)
                        break

            # Limpa slots sobrando (quando há menos questões do que slots no template)
            n_preenchidas = len(questoes_dissertativas)
            for idx_extra in range(9 + n_preenchidas, 9 + _N_SLOTS_DIS):
                ph_extra = f"{{{{Questão {idx_extra} aqui}}}}"
                for paragraph in document.paragraphs:
                    if ph_extra in paragraph.text:
                        paragraph.clear()   # remove numeração e placeholder
                        break

        # Separa as discursivas para seção de 1 coluna (largura total)
        if tipo_avaliacao == "R" and paragrafos_dissertativos:
            _separar_secao_dissertativas(document)
            # Injeta cabeçalho "Questões Discursivas [= Y pts]" sempre
            _injetar_titulo_dissertativas(document, paragrafos_dissertativos[0], pontos_dis)

        # Injeta linhas de resposta após cada questão discursiva (ou após a imagem, se houver)
        if paragrafos_ultimos_dis and linhas_por_questao > 0:
            for dp in paragrafos_ultimos_dis:
                _injetar_linhas_apos(dp, linhas_por_questao)

        # Substitui o símbolo da versão no rodapé
        for paragraph in all_paragraphs(document):
            replace_text_in_paragraph_runs(paragraph, "{{simbolo aqui}}", simbolo_rodape)

        # Substitui semestre e ano
        agora    = datetime.now()
        semestre = "1" if agora.month <= 6 else "2"
        ano      = str(agora.year)
        for paragraph in all_paragraphs(document):
            replace_text_in_paragraph_runs(paragraph, "{{Sem aqui}}", semestre)
        for paragraph in all_paragraphs(document):
            replace_text_in_paragraph_runs(paragraph, "{{Ano aqui}}", ano)

        # Substitui professor e disciplina
        for paragraph in all_paragraphs(document):
            replace_text_in_paragraph_runs(paragraph, "{{professor aqui}}", professor)
        for paragraph in all_paragraphs(document):
            replace_text_in_paragraph_runs(paragraph, "{{disciplina aqui}}", disciplina)

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except Exception as e:
        st.error(f"Erro ao gerar prova {identificador_prova}: {e}")
        return None


def gera_gabarito_str(gabaritos_por_prova, tipo_avaliacao,
                      gabaritos_dissertativas_por_prova=None):
    """
    Gera o conteúdo do gabarito geral como string UTF-8.
    gabaritos_dissertativas_por_prova: dict {nome_prova: [(enunciado, resposta), ...]}
    """
    nome_completo = "Regimental" if tipo_avaliacao == "R" else "Final"
    linhas = [f"--- GABARITO GERAL - AVALIAÇÃO {nome_completo.upper()} ---\n"]
    for nome_prova, gabarito_texto in gabaritos_por_prova.items():
        linhas.append(f"PROVA {nome_prova}\n{gabarito_texto}")
        # Gabarito das dissertativas desta versão
        if gabaritos_dissertativas_por_prova:
            dis_desta_prova = gabaritos_dissertativas_por_prova.get(nome_prova, [])
            if dis_desta_prova:
                linhas.append("\nGABARITO — QUESTÕES DISSERTATIVAS")
                for i, (enunciado, resposta) in enumerate(dis_desta_prova, start=9):
                    linhas.append(f"\nQuestão {i}:")
                    linhas.append(f"  Enunciado: {enunciado}")
                    linhas.append(f"  Resposta esperada: {resposta}")
        linhas.append(f"\n{'=' * 40}\n")
    return "\n".join(linhas)


# --- Interface Streamlit ---

def _ui_importar_ia():
    """
    Expander de importação de documentos não estruturados via Claude API.
    Armazena as questões extraídas em st.session_state e limpa a chave API
    imediatamente após a chamada.
    """
    with st.expander("Importar documento não estruturado via IA", expanded=False):
        if not ANTHROPIC_DISPONIVEL:
            st.warning(
                "Pacote `anthropic` não instalado.  \n"
                "Execute no terminal: `pip install anthropic`"
            )
            return

        # Contador usado para trocar a key do campo de senha e limpá-lo após o uso.
        # Streamlit não permite setar o valor de um widget pela key na mesma execução
        # em que ele foi renderizado — mudar a key força um widget novo e vazio.
        if "api_key_run_id" not in st.session_state:
            st.session_state["api_key_run_id"] = 0

        # Mensagem de sucesso diferida (exibida no rerun pós-processamento)
        _msg_sucesso = st.session_state.get("ia_msg_sucesso")
        if _msg_sucesso:
            del st.session_state["ia_msg_sucesso"]
            st.success(_msg_sucesso)

        st.caption(
            "Envie um documento com questões em qualquer formato. "
            "A IA identifica e estrutura automaticamente as questões.  \n"
            "Formatos aceitos: **.docx** · **.xlsx** · **.txt** · **.gs** (Apps Script) · **.json** (Google Forms API/export)."
        )

        col_key, col_openai, col_tipo = st.columns([2, 2, 1])
        with col_key:
            api_key_input = st.text_input(
                "Chave API Anthropic",
                type="password",
                placeholder="sk-ant-...",
                key=f"api_key_widget_{st.session_state['api_key_run_id']}",
                help="Usada somente nesta chamada. Apagada automaticamente após o processamento.",
            )
        with col_openai:
            openai_key_input = st.text_input(
                "Chave API OpenAI (DALL-E)",
                type="password",
                placeholder="sk-...",
                key=f"openai_key_widget_{st.session_state['api_key_run_id']}",
                help="Opcional. Necessária para sugerir e gerar imagens com DALL-E 3.",
            )

        # Persiste ambas as chaves na sessão enquanto o usuário as digitar
        if api_key_input.strip():
            st.session_state["anthropic_key_salva"] = api_key_input.strip()
        if openai_key_input.strip():
            st.session_state["openai_key_salva"] = openai_key_input.strip()
        with col_tipo:
            tipo_extracao = st.selectbox(
                "Extrair",
                ["Ambas (objetivas + discursivas)", "Apenas objetivas", "Apenas discursivas"],
                key="tipo_extracao_ia",
            )

        arquivo_ia = st.file_uploader(
            "Carregar documento (.docx, .xlsx, .txt, .gs, .json)",
            type=["docx", "xlsx", "txt", "gs", "json"],
            key="arquivo_ia",
        )

        # Botões de ação — sempre renderizados; "Regenerar" desabilitado sem questões
        _tem_questoes = bool(
            st.session_state.get("ia_obj") or st.session_state.get("ia_dis")
        )
        col_proc, col_regen = st.columns(2)
        with col_proc:
            processar = st.button("Processar com IA", key="btn_processar_ia")
        with col_regen:
            _regenerar = st.button(
                "Regenerar nao confirmadas",
                key="btn_regenerar_ia",
                disabled=not _tem_questoes,
                help="Reenvia apenas as questoes sem aprovacao para a IA gerar novas versoes.",
            )

        # --- Processamento inicial ---
        if processar:
            if not api_key_input.strip():
                st.error("Insira a chave API antes de processar.")
            elif arquivo_ia is None:
                st.error("Carregue um documento antes de processar.")
            else:
                with st.spinner("Extraindo objetivas e criando discursivas — podem ser 2 chamadas à API..."):
                    texto, erro_texto = extrair_texto_arquivo(arquivo_ia)
                    if erro_texto:
                        st.error(f"Erro ao ler arquivo: {erro_texto}")
                    else:
                        tipo_map = {
                            "Ambas (objetivas + discursivas)": "ambas",
                            "Apenas objetivas": "objetivas",
                            "Apenas discursivas": "dissertativas",
                        }
                        resultado, erro_api = processar_com_claude(
                            texto, api_key_input.strip(), tipo_map[tipo_extracao],
                        )
                        if erro_api:
                            st.error(f"Erro: {erro_api}")
                        else:
                            obj, dis, gab_dis = _converter_resultado_ia(resultado)
                            st.session_state["ia_obj"]        = obj
                            st.session_state["ia_dis"]        = dis
                            st.session_state["ia_dis_gab"]    = gab_dis
                            st.session_state["ia_obj_ok"]     = [True] * len(obj)
                            st.session_state["ia_dis_ok"]     = [True] * len(dis)
                            st.session_state["ia_texto_doc"]  = texto
                            st.session_state["ia_confirmadas"] = False
                            st.session_state["ia_msg_sucesso"] = (
                                f"Extracao concluida: **{len(obj)}** objetiva(s) · "
                                f"**{len(dis)}** discursiva(s).  \n"
                                "Revise abaixo, aprove cada questao e clique em **Usar questoes aprovadas**."
                            )
                            st.session_state["api_key_run_id"] += 1
                            st.rerun()

        # --- Regeneração seletiva ---
        if _regenerar:
            obj_all  = st.session_state.get("ia_obj", [])
            dis_all  = st.session_state.get("ia_dis", [])
            obj_ok   = st.session_state.get("ia_obj_ok", [True] * len(obj_all))
            dis_ok   = st.session_state.get("ia_dis_ok", [True] * len(dis_all))
            texto_doc = st.session_state.get("ia_texto_doc", "")

            obj_conf = [q for q, ok in zip(obj_all, obj_ok) if ok]
            dis_conf = [q for q, ok in zip(dis_all, dis_ok) if ok]
            n_obj_falt = sum(1 for ok in obj_ok if not ok)
            n_dis_falt = sum(1 for ok in dis_ok if not ok)

            if not api_key_input.strip():
                st.error("Insira a chave API para regenerar.")
            elif n_obj_falt == 0 and n_dis_falt == 0:
                st.info("Todas as questoes ja estao aprovadas — nada a regenerar.")
            elif not texto_doc:
                st.error("Texto do documento nao disponivel. Reprocesse o arquivo.")
            else:
                tipo_map = {
                    "Ambas (objetivas + discursivas)": "ambas",
                    "Apenas objetivas": "objetivas",
                    "Apenas discursivas": "dissertativas",
                }
                with st.spinner(f"Regenerando {n_obj_falt} objetiva(s) e {n_dis_falt} discursiva(s)..."):
                    resultado, erro_api = regenerar_nao_confirmadas(
                        texto_doc, api_key_input.strip(), tipo_map[tipo_extracao],
                        obj_conf, dis_conf, n_obj_falt, n_dis_falt,
                    )
                    if erro_api:
                        st.error(f"Erro: {erro_api}")
                    else:
                        novas_obj, novas_dis, novas_gab = _converter_resultado_ia(resultado)
                        gab_all = list(st.session_state.get("ia_dis_gab", [""] * len(dis_all)))

                        # Substitui slots não aprovados pelas novas questões
                        obj_final, ok_final = list(obj_all), list(obj_ok)
                        slot_obj = 0
                        for i in range(len(obj_final)):
                            if not ok_final[i] and slot_obj < len(novas_obj):
                                obj_final[i] = novas_obj[slot_obj]
                                ok_final[i]  = False
                                slot_obj += 1

                        dis_final, dok_final, gab_final = list(dis_all), list(dis_ok), gab_all
                        slot_dis = 0
                        for i in range(len(dis_final)):
                            if not dok_final[i] and slot_dis < len(novas_dis):
                                dis_final[i] = novas_dis[slot_dis]
                                gab_final[i] = novas_gab[slot_dis] if slot_dis < len(novas_gab) else ""
                                dok_final[i] = False
                                slot_dis += 1

                        st.session_state["ia_obj"]     = obj_final
                        st.session_state["ia_dis"]     = dis_final
                        st.session_state["ia_dis_gab"] = gab_final
                        st.session_state["ia_obj_ok"]  = ok_final
                        st.session_state["ia_dis_ok"]  = dok_final
                        st.session_state["api_key_run_id"] += 1
                        st.session_state["ia_msg_sucesso"] = (
                            f"Regeneracao concluida: {slot_obj} objetiva(s) e "
                            f"{slot_dis} discursiva(s) substituidas. Revise abaixo."
                        )
                        st.rerun()

        # --- Painel de revisao por questão ---
        obj_prev = st.session_state.get("ia_obj", [])
        dis_prev = st.session_state.get("ia_dis", [])

        if obj_prev or dis_prev:
            # Garante que os status existem e têm o tamanho correto
            if "ia_obj_ok" not in st.session_state or len(st.session_state["ia_obj_ok"]) != len(obj_prev):
                st.session_state["ia_obj_ok"] = [True] * len(obj_prev)
            if "ia_dis_ok" not in st.session_state or len(st.session_state["ia_dis_ok"]) != len(dis_prev):
                st.session_state["ia_dis_ok"] = [True] * len(dis_prev)

            obj_ok = st.session_state["ia_obj_ok"]
            dis_ok = st.session_state["ia_dis_ok"]

            n_ok = sum(obj_ok) + sum(dis_ok)
            n_total = len(obj_prev) + len(dis_prev)
            st.markdown(f"**Revisao das questoes extraidas** — {n_ok}/{n_total} aprovadas")

            # Objetivas
            if obj_prev:
                st.markdown("**Questoes Objetivas**")
                for i, q in enumerate(obj_prev):
                    icone = "OK" if obj_ok[i] else "X"
                    cor   = "green" if obj_ok[i] else "red"
                    with st.container(border=True):
                        c_icon, c_texto, c_btn = st.columns([0.08, 0.80, 0.12])
                        with c_icon:
                            st.markdown(
                                f"<span style='font-size:1.4rem;color:{cor}'>"
                                f"{'✅' if obj_ok[i] else '❌'}</span>",
                                unsafe_allow_html=True,
                            )
                        with c_texto:
                            st.markdown(f"**Q{i+1}.** {q[0]}")
                            st.caption(
                                f"**Correta:** {q[1]}  \n"
                                f"**Distratores:** {' / '.join(q[3:7])}"
                            )
                        with c_btn:
                            label = "Reprovar" if obj_ok[i] else "Aprovar"
                            if st.button(label, key=f"tog_obj_{i}"):
                                st.session_state["ia_obj_ok"][i] = not obj_ok[i]
                                st.rerun()

            # Dissertativas
            if dis_prev:
                gab_prev = st.session_state.get("ia_dis_gab", [""] * len(dis_prev))
                if len(gab_prev) != len(dis_prev):
                    gab_prev = [""] * len(dis_prev)

                st.markdown("**Questoes Discursivas**")
                for i, q in enumerate(dis_prev):
                    with st.container(border=True):
                        c_icon, c_texto, c_btn = st.columns([0.08, 0.80, 0.12])
                        with c_icon:
                            st.markdown(
                                f"<span style='font-size:1.4rem;color:{'green' if dis_ok[i] else 'red'}'>"
                                f"{'✅' if dis_ok[i] else '❌'}</span>",
                                unsafe_allow_html=True,
                            )
                        with c_texto:
                            st.markdown(f"**D{i+1}.** {q}")
                            if gab_prev[i]:
                                st.caption(f"**Resposta esperada:** {gab_prev[i]}")
                        with c_btn:
                            label = "Reprovar" if dis_ok[i] else "Aprovar"
                            if st.button(label, key=f"tog_dis_{i}"):
                                st.session_state["ia_dis_ok"][i] = not dis_ok[i]
                                st.rerun()

            # --- Painel de imagens DALL-E ---
            st.divider()
            _openai_key_atual   = st.session_state.get("openai_key_salva", "")
            _anthropic_key_img  = st.session_state.get("anthropic_key_salva", "")

            with st.expander("Imagens com DALL-E 3 (opcional)", expanded=bool(st.session_state.get("ia_imgs_preview") or st.session_state.get("ia_imgs_prompts"))):
                st.caption(
                    "**Passo 1** — Claude sugere quais questões se beneficiam de imagem e gera os prompts.  \n"
                    "**Passo 2** — DALL-E 3 gera cada imagem. Você pode regenerar sem repetir o Passo 1."
                )

                if not OPENAI_DISPONIVEL:
                    st.warning("Pacote `openai` não instalado. Execute: `pip install openai`")
                else:
                    _obj_aprov_img = [q for q, ok in zip(obj_prev, obj_ok) if ok]
                    _dis_aprov_img = [q for q, ok in zip(dis_prev, dis_ok) if ok]
                    _n_aprov_img   = len(_obj_aprov_img) + len(_dis_aprov_img)
                    _prompts_salvos = st.session_state.get("ia_imgs_prompts", {})
                    _preview_atual  = st.session_state.get("ia_imgs_preview", {})
                    _erros_salvos   = st.session_state.get("ia_imgs_erros", {})   # {enunc: True}
                    _n_com_erro     = len(_erros_salvos)

                    # --- Passo 1: sugerir prompts com Claude ---
                    col_p1, col_p2 = st.columns(2)
                    with col_p1:
                        _btn_sugerir = st.button(
                            f"Passo 1 — Sugerir prompts ({_n_aprov_img} questões)",
                            key="btn_sugerir_prompts",
                            help="Requer chave Anthropic. Só precisa rodar uma vez.",
                            disabled=not _anthropic_key_img,
                        )
                    with col_p2:
                        _tem_prompts  = bool(_prompts_salvos)
                        _n_pendentes  = sum(1 for e in _prompts_salvos if e not in _preview_atual or e in _erros_salvos)
                        _label_p2 = (
                            f"Passo 2 — Regenerar com falha ({_n_com_erro})" if _n_com_erro
                            else f"Passo 2 — Gerar imagens ({len(_prompts_salvos)} prompts)"
                        )
                        _btn_gerar = st.button(
                            _label_p2,
                            key="btn_gerar_dalle",
                            help="Usa os prompts já gerados no Passo 1. Requer chave OpenAI.",
                            disabled=not (_tem_prompts and _openai_key_atual),
                        )

                    if not _anthropic_key_img and not _prompts_salvos:
                        st.info("Insira a Chave API Anthropic e clique em Passo 1 para começar.")
                    if not _openai_key_atual:
                        st.info("Insira a Chave API OpenAI (DALL-E) acima para habilitar o Passo 2.")

                    # Executa Passo 1
                    if _btn_sugerir:
                        with st.spinner("Claude gerando prompts de imagem para cada questão..."):
                            try:
                                _pm, _err_p = _gerar_prompts_imagem_claude(
                                    _obj_aprov_img, _dis_aprov_img, _anthropic_key_img
                                )
                            except Exception as _ex_p1:
                                _pm, _err_p = None, str(_ex_p1)
                        if _err_p:
                            st.error(f"Erro ao gerar prompts: {_err_p}")
                        else:
                            st.session_state["ia_imgs_prompts"] = _pm
                            st.session_state["ia_imgs_erros"]   = {}  # limpa erros anteriores
                            n_pm = len(_pm)
                            st.success(f"{n_pm} prompt(s) gerado(s). Clique em **Passo 2** para criar as imagens.")
                            st.rerun()

                    # Executa Passo 2 (geração DALL-E) — só processa pendentes/com erro
                    if _btn_gerar and _prompts_salvos and _openai_key_atual:
                        _preview_novo  = dict(_preview_atual)
                        _erros_novo    = {}
                        _erros_msg_novo = {}
                        _ok_imgs       = dict(st.session_state.get("ia_imgs_ok", {}))
                        _pendentes = {
                            e: p for e, p in _prompts_salvos.items()
                            if e not in _preview_atual or e in _erros_salvos
                        }
                        _n_pend = len(_pendentes)
                        if _n_pend == 0:
                            st.info("Todas as imagens já foram geradas com sucesso.")
                        else:
                            _prog = st.progress(0, text=f"Gerando {_n_pend} imagem(ns) com DALL-E 3...")
                            for _pi, (_enunc, _dalle_p) in enumerate(_pendentes.items()):
                                _prog.progress((_pi + 1) / _n_pend,
                                               text=f"Gerando imagem {_pi+1}/{_n_pend}...")
                                _img_b, _err_i = _gerar_imagem_dalle(_dalle_p, _openai_key_atual)
                                if _img_b:
                                    _preview_novo[_enunc] = _img_b
                                    _ok_imgs[_enunc] = _ok_imgs.get(_enunc, True)
                                else:
                                    _erros_novo[_enunc] = True
                                    _erros_msg_novo[_enunc] = str(_err_i)
                            _prog.empty()
                            st.session_state["ia_imgs_preview"]  = _preview_novo
                            st.session_state["ia_imgs_ok"]       = _ok_imgs
                            st.session_state["ia_imgs_erros"]    = _erros_novo
                            st.session_state["ia_imgs_erros_msg"] = _erros_msg_novo
                            st.rerun()

                    # Exibe erros persistentes
                    _erros_msg_salvos = st.session_state.get("ia_imgs_erros_msg", {})
                    if _erros_msg_salvos:
                        for _enunc_err, _msg_err in _erros_msg_salvos.items():
                            st.error(f"**Falha ao gerar imagem:** {_msg_err}  \n`Questão: {_enunc_err[:120]}`")

                    # Galeria de prévia
                    if _preview_atual:
                        _ok_imgs_atual = st.session_state.get("ia_imgs_ok", {})
                        n_geradas = len(_preview_atual)
                        n_erros   = len(_erros_salvos)
                        st.markdown(f"**{n_geradas} gerada(s)** · {n_erros} com falha — aprove ou rejeite cada imagem:")
                        for _enunc_img, _img_b in _preview_atual.items():
                            with st.container(border=True):
                                c_img, c_info = st.columns([0.45, 0.55])
                                with c_img:
                                    st.image(_img_b, use_container_width=True)
                                with c_info:
                                    _aprovada = _ok_imgs_atual.get(_enunc_img, True)
                                    st.markdown(
                                        f"<span style='color:{'green' if _aprovada else 'red'};font-size:1.2rem'>"
                                        f"{'✅ Aprovada' if _aprovada else '❌ Rejeitada'}</span>",
                                        unsafe_allow_html=True,
                                    )
                                    st.caption(_enunc_img[:200])
                                    _label_img = "Rejeitar" if _aprovada else "Aprovar"
                                    if st.button(_label_img, key=f"tog_img_{hash(_enunc_img) % 99999}"):
                                        st.session_state["ia_imgs_ok"][_enunc_img] = not _aprovada
                                        st.rerun()

            st.divider()
            col_usar, col_limpar = st.columns(2)
            with col_usar:
                n_obj_aprov = sum(obj_ok)
                n_dis_aprov = sum(dis_ok)
                if st.button(
                    f"Usar questoes aprovadas ({n_obj_aprov} obj · {n_dis_aprov} dis)",
                    type="primary", key="btn_usar_ia"
                ):
                    gab_all = st.session_state.get("ia_dis_gab", [""] * len(dis_prev))
                    # Filtra apenas as aprovadas (questões e gabaritos alinhados)
                    st.session_state["ia_obj"]     = [q for q, ok in zip(obj_prev, obj_ok) if ok]
                    st.session_state["ia_dis"]     = [q for q, ok in zip(dis_prev, dis_ok) if ok]
                    st.session_state["ia_dis_gab"] = [g for g, ok in zip(gab_all, dis_ok) if ok]
                    st.session_state["ia_obj_ok"]  = [True] * len(st.session_state["ia_obj"])
                    st.session_state["ia_dis_ok"]  = [True] * len(st.session_state["ia_dis"])
                    st.session_state["ia_confirmadas"] = True
                    # Salva imagens aprovadas: {enunciado: bytes}
                    _prev = st.session_state.get("ia_imgs_preview", {})
                    _ok_i = st.session_state.get("ia_imgs_ok", {})
                    st.session_state["ia_imgs"] = {
                        enunc: b for enunc, b in _prev.items() if _ok_i.get(enunc, True)
                    }
                    st.rerun()
            with col_limpar:
                if st.button("Descartar tudo", key="btn_limpar_ia"):
                    for k in ("ia_obj", "ia_dis", "ia_dis_gab", "ia_obj_ok",
                              "ia_dis_ok", "ia_confirmadas", "ia_texto_doc",
                              "ia_imgs", "ia_imgs_preview", "ia_imgs_ok",
                              "ia_imgs_prompts", "ia_imgs_erros",
                              "openai_key_salva"):
                        if k in st.session_state:
                            del st.session_state[k]
                    st.rerun()

        if st.session_state.get("ia_confirmadas"):
            n_obj = len(st.session_state.get("ia_obj", []))
            n_dis = len(st.session_state.get("ia_dis", []))
            st.success(
                f"Questoes ativas via IA: **{n_obj}** objetiva(s) · **{n_dis}** discursiva(s). "
                "Pronto para gerar a prova."
            )
            # Dica de variedade: pool pequeno = mesmas questões em todas as versões
            avisos = []
            if n_obj <= 8:
                avisos.append(
                    f"Objetivas: pool de {n_obj} questão(ões) — todas as versões usarão as mesmas."
                    " Processe novamente para obter questões diferentes."
                )
            if n_dis <= 2:
                avisos.append(
                    f"Discursivas: pool de {n_dis} questão(ões) — sem variedade entre versões."
                    " Processe novamente para ampliar o pool."
                )
            if avisos:
                st.info(
                    "**Dica de variedade entre versões:** " + " | ".join(avisos)
                )


_EVIDENCIAS = os.path.join(BASE_DIR, "Evidencias_manual_ProvaRegimental")


def _img(nome):
    """Exibe uma imagem da pasta de evidências, se existir."""
    caminho = os.path.join(_EVIDENCIAS, nome)
    if os.path.exists(caminho):
        st.image(caminho)


def _ui_manual():
    st.header("Manual do Usuário")
    st.caption("Versão 3.6.3 · Prof.Me. Cid R. Andrade · Co-Autor: Prof.Me. Rafael Cotrin (v3.4.0+)")
    st.divider()

    # 1. Introdução
    st.subheader("1. Introdução")
    st.markdown("""
O **Gerador de Provas Unicid** automatiza a criação de avaliações institucionais no padrão Unicid,
produzindo múltiplas versões embaralhadas de uma mesma prova para garantir integridade acadêmica.

**Principais capacidades:**
- Geração de provas **AR** (Regimental) e **AF** (Final) em DOCX e PDF.
- Dois fluxos de entrada: **planilha XLSX** (método tradicional) ou **importação via IA** (Claude AI).
- Até **8 versões** (A–H) com embaralhamento automático de questões e alternativas.
- Exportação em **ZIP** com DOCXs, PDFs opcionais e gabarito em texto.
""")

    # 2. Requisitos
    st.subheader("2. Requisitos")
    st.markdown("""
| Requisito | Observação |
|---|---|
| Navegador atualizado | Chrome, Edge, Firefox ou Safari |
| Conexão com internet | Necessária para acessar o sistema e para o fluxo via IA |
| LibreOffice ou Microsoft Word | Necessário **apenas** para geração de PDF — LibreOffice (Linux/Streamlit Cloud) ou Word (Windows/Mac) |
| Chave da API Anthropic | Necessária **apenas** para o fluxo de importação via IA |

> O sistema roda integralmente na nuvem — nenhum software adicional precisa ser instalado pelo usuário.
""")

    # 3. Como acessar
    st.subheader("3. Como Acessar")
    st.markdown("""
Acesse pelo navegador, sem instalação:

**[https://prova-unicid.streamlit.app/](https://prova-unicid.streamlit.app/)**

> Compatível com qualquer navegador moderno (Chrome, Edge, Firefox, Safari).
> Não é necessário instalar nenhum software para usar o sistema.
""")

    # 4. Visão geral
    st.subheader("4. Visão Geral da Interface")
    _img("Carregamento de questões com ou sem IA.png")
    st.markdown("""
A interface divide-se em dois painéis lado a lado:

- **Painel esquerdo:** configurações da prova (tipo, gabarito, versões, professor, disciplina, pontuação).
- **Painel direito:** carregamento de questões (XLSX ou via IA).

Acima dos painéis: o expander **"Importar documento não estruturado via IA"** (opcional).
""")

    # 5. Fluxo XLSX
    st.subheader("5. Fluxo 1 — Questões via Planilha XLSX")
    st.markdown("""
Indicado quando o professor já possui as questões organizadas em planilha.

**Passo a passo:**
1. Selecione o **Tipo de avaliação** (AR ou AF).
2. Selecione o **Tipo de gabarito** (Padrão ou Zipgrade).
3. Ajuste o **Número de versões**.
4. Preencha **Professor** e **Disciplina**.
5. Faça upload da planilha de **questões objetivas** (`.xlsx`).
6. Se AR: faça upload da planilha de **questões discursivas** (`.xlsx`).
7. Configure pontuação, linhas de resposta e clique em **Gerar Provas**.

**Formato — Questões Objetivas:**

| Coluna | Conteúdo |
|---|---|
| A | Enunciado |
| B | Resposta correta |
| C–F | Distratores (4 alternativas incorretas) |

**Formato — Questões Discursivas:**

| Coluna | Conteúdo |
|---|---|
| A | Enunciado |

**Formatação de texto nas células:**

| Marcação | Resultado |
|---|---|
| `*texto*` | *itálico* |
| `**texto**` | **negrito** |
| `***texto***` | ***negrito e itálico*** |
""")

    # 6. Fluxo IA
    st.subheader("6. Fluxo 2 — Importação via IA")
    _img("Esrutura com IA.png")
    st.markdown("""
Utiliza a API do **Claude AI (Anthropic)** para extrair e estruturar questões de documentos não estruturados.

> **Custo estimado:** ~US$ 0,10 por processamento (varia com o volume de texto).

**Formatos aceitos:** `.docx`, `.xlsx`, `.txt`, `.gs` (Google Apps Script), `.json` (Google Forms)

**Passo a passo:**
1. Insira sua **chave da API** (começa com `sk-ant-...`).
2. Selecione o que extrair: **Ambas**, **Objetivas** ou **Discursivas**.
3. Faça upload do documento de origem.
4. Clique em **Processar com IA**.
""")
    st.markdown("**Revisão das questões geradas:**")
    _img("Saída do modelo para avaliacao do professor.png")
    st.markdown("""
Cada questão extraída é exibida para revisão:
- **✅ Aprovar** — inclui no pool ativo.
- **❌ Rejeitar** — descarta da sessão.

Use **Regenerar não confirmadas** para pedir uma nova tentativa só para as rejeitadas.
""")
    st.markdown("**Confirmação:**")
    _img("Confirmação_dos dados.png")
    st.markdown("""
- **Usar questões aprovadas** — confirma o pool e habilita a geração.
- **Descartar tudo** — remove tudo e permite recomeçar.

> A chave da API **não é salva** — é descartada após cada processamento.
""")

    # 7. Configurações
    st.subheader("7. Configurações da Prova")
    _img("Configurações Finais para gerar as provas.png")
    st.markdown("""
| Configuração | Descrição |
|---|---|
| **Tipo de avaliação** | AR: 8 obj + 2 ou 3 dis · AF: 20 obj |
| **Tipo de gabarito** | Padrão ou Zipgrade (correção por leitura óptica) |
| **Número de versões** | 1 a 8 (letras A–H) |
| **Professor / Disciplina** | Inseridos no cabeçalho de todas as versões |
| **Questões dissertativas** | Radio 2 ou 3 — calcula automaticamente o valor por questão |
| **Linhas de resposta** | Slider 4–20 linhas por questão dissertativa |
| **Incluir PDF** | Requer LibreOffice (Linux/Streamlit Cloud) ou Microsoft Word (Windows/Mac) |

**Pontuação calculada automaticamente (AR):**
- Objetivas: 8 × 0,25 pt = **2,00 pts** (fixo)
- Discursivas: 3,00 pts ÷ nº de questões = **1,00 pt** (3 dis) ou **1,50 pt** (2 dis)
""")

    # 8. Geração e download
    st.subheader("8. Geração e Download")
    st.markdown("""
Após configurar, clique em **Gerar Provas**. O ZIP baixado contém:

- `prova_A.docx`, `prova_B.docx`, ... (um por versão)
- `prova_A.pdf`, `prova_B.pdf`, ... (se PDF habilitado)
- `Gabarito_Geral_*.txt` — gabarito de todas as versões

**Estrutura do DOCX:**
- Cabeçalho com dados institucionais e identificador da versão.
- Questões objetivas em **2 colunas**.
- Quebra para nova página, questões discursivas em **1 coluna** com linhas de resposta.
""")

    # 9. Dicas
    st.subheader("9. Dicas de Uso")
    st.markdown("""
- Inclua **o dobro de questões** do necessário na planilha para maximizar a variação entre versões.
- Ao usar a IA, prefira documentos com enunciados **numerados e bem delimitados**.
- **Revise o gabarito** após cada geração — o embaralhamento muda a ordem das respostas.
- Salve sua chave da API em um **gerenciador de senhas**; o sistema não a armazena.
- Faça um **teste com 1 versão** antes de gerar o conjunto completo para validar o layout.
""")

    # 10. Limitações
    st.subheader("10. Limitações Conhecidas")
    st.markdown("""
- Geração de PDF requer LibreOffice (Linux/Streamlit Cloud) ou Microsoft Word (Windows/Mac).
- O fluxo via IA exige internet durante o processamento.
- A qualidade da extração depende da organização do documento de origem.
- O sistema não valida academicamente as questões geradas pela IA — revisão é indispensável.
- Máximo de 8 versões por geração.
""")


def _ui_faq():
    st.header("FAQ — Perguntas Frequentes")
    st.divider()

    perguntas = [
        ("Instalação e Requisitos",
         "Preciso do Microsoft Word ou LibreOffice instalado?",
         """Para gerar as provas em **DOCX**, nenhum deles é necessário.
         Se marcar **"Incluir PDF no download"**, o sistema usa **LibreOffice** (Linux/Streamlit Cloud)
         ou **Microsoft Word** (Windows/Mac) para a conversão.
         Sem o conversor disponível, o DOCX é gerado normalmente; apenas o PDF fica indisponível."""),

        ("Instalação e Requisitos",
         "O sistema funciona sem internet?",
         """**Não.** O sistema é hospedado em nuvem e exige conexão para ser acessado.
         Dentro do sistema, o fluxo via **XLSX não depende de serviços externos** —
         toda a geração é feita no servidor. Já o fluxo via **IA exige conexão adicional**
         com a API do Claude (Anthropic) durante o processamento do documento."""),

        ("Instalação e Requisitos",
         "Quais arquivos de template são necessários?",
         """`Modelo_AR.docx`, `Modelo_ARZ.docx`, `Modelo_AF.docx`, `Modelo_AFZ.docx` —
         todos na mesma pasta do aplicativo. Sem eles, o sistema não consegue gerar as provas.
         Mantenha backups desses arquivos."""),

        ("Planilha XLSX",
         "Qual o formato correto da planilha de questões objetivas?",
         """Cada linha = uma questão. Colunas: **A** = enunciado · **B** = resposta correta ·
         **C–F** = distratores (4 alternativas incorretas). Sem cabeçalho obrigatório.
         O sistema embaralha as alternativas automaticamente em cada versão."""),

        ("Planilha XLSX",
         "Posso reutilizar planilhas de semestres anteriores?",
         """Sim. As planilhas XLSX não têm prazo de validade.
         Recomenda-se manter um arquivo por disciplina e acrescentar questões ao longo do tempo
         para ampliar o pool disponível."""),

        ("Importação via IA",
         "A IA sempre gera as mesmas questões?",
         """Não. A IA usa amostragem probabilística — cada processamento pode produzir variações
         mesmo a partir do mesmo documento. Além disso, o sistema injeta uma instrução de foco
         diferente a cada chamada para ampliar a diversidade. Use **Regenerar não confirmadas**
         para obter novas tentativas."""),

        ("Importação via IA",
         "Quais formatos são aceitos para importação via IA?",
         """`.docx`, `.xlsx` (sem formato padrão), `.txt`, `.gs` (Google Apps Script),
         `.json` (Google Forms API/export). Qualquer documento elaborado pelo professor
         pode ser utilizado."""),

        ("Importação via IA",
         "Posso usar Google Forms como fonte?",
         """Sim, de três formas:
         1. Exporte o formulário como `.json` via Google Forms API e importe diretamente.
         2. Exporte o script do formulário em `.gs` (Google Apps Script).
         3. Copie as questões para um `.docx` ou `.txt` e importe."""),

        ("Importação via IA",
         "A chave da API fica salva no sistema?",
         """Não. A chave é usada apenas durante a sessão ativa e descartada logo após
         o processamento. Ela não é gravada em disco nem enviada a nenhum servidor
         além da API da Anthropic. Guarde-a em um gerenciador de senhas."""),

        ("Importação via IA",
         "Qual o custo de uso da IA?",
         """~US$ 0,10 por processamento (varia com o volume de texto e o modelo).
         Cada clique em **Processar com IA** ou **Regenerar não confirmadas** gera
         uma nova chamada com custo adicional. Consulte [anthropic.com/pricing](https://anthropic.com/pricing)."""),

        ("Importação via IA",
         "O que fazer se a IA não gerou questões discursivas?",
         """1. Verifique se **Extrair** está em "Ambas" ou "Discursivas".
         2. Tente **Regenerar não confirmadas**.
         3. Forneça um documento mais estruturado (enunciados numerados ou separados).
         A IA **cria** as discursivas a partir dos temas — o documento de origem não precisa tê-las."""),

        ("Importação via IA",
         "Posso combinar IA para discursivas e XLSX para objetivas?",
         """Sim. Selecione **"Discursivas"** no dropdown da IA, aprove as questões
         e faça upload da planilha de objetivas normalmente. O sistema combina os dois."""),

        ("Configuração e Geração",
         "Quantas versões posso gerar?",
         """Até **8 versões** (A–H) por geração. Para turmas maiores, repita o processo
         com um pool diferente de questões para obter versões distintas."""),

        ("Configuração e Geração",
         "O que é Zipgrade?",
         """Aplicativo (iOS/Android) que corrige provas por leitura óptica, como um gabarito
         de bolinha fotografado pelo celular. Ao selecionar **Zipgrade**, o sistema usa o template
         compatível que inclui o cartão-resposta."""),

        ("Configuração e Geração",
         "Como garantir variedade entre as versões?",
         """O pool precisa ser **maior que o mínimo necessário**:
         - Via XLSX: inclua mais linhas do que o necessário.
         - Via IA: aprove mais questões do que o mínimo.

         **Exemplo:** 16 objetivas aprovadas para AR (que precisa de 8) = cada versão pode
         receber 8 questões completamente diferentes. O sistema exibe dica de variedade
         quando o pool é pequeno."""),

        ("Configuração e Geração",
         "O que significa 'pool de questões'?",
         """É o conjunto total de questões disponíveis para sorteio. Para cada versão,
         o sistema seleciona aleatoriamente o número necessário a partir deste pool.
         Quanto maior o pool, maior a variedade entre versões."""),

        ("Configuração e Geração",
         "Como funciona o gabarito gerado?",
         """O arquivo `.txt` no ZIP lista a sequência de respostas corretas para cada versão,
         já considerando o embaralhamento aplicado. A resposta certa da Questão 1 na Versão A
         pode ser diferente da Versão B — o gabarito reflete isso. Sempre confira antes da prova."""),

        ("Problemas",
         "O sistema trava ou apresenta erro. O que fazer?",
         """Causas mais comuns:
         1. **Templates ausentes** — verifique os quatro `Modelo_*.docx` na pasta.
         2. **Planilha com formato incorreto** — confira as colunas (A=enunciado, B=correta, C-F=distratores).
         3. **Questões insuficientes** — AR exige ≥8 obj + ≥2 disc; AF exige ≥20 obj.
         4. **Erro de API** — verifique a chave e o saldo na conta Anthropic.
         5. **Recarregue a página** no navegador (`F5`) e tente novamente."""),

        ("Problemas",
         "O PDF gerado é diferente do DOCX?",
         """Não. O PDF é uma conversão fiel do DOCX — mesmo conteúdo e layout.
         É fornecido apenas como alternativa de formato para impressão ou distribuição."""),

        ("Problemas",
         "Posso gerar Prova Final (AF) usando a IA?",
         """Sim. Selecione **AF**, use o dropdown da IA em "Objetivas" (ou "Ambas"),
         aprove ao menos 20 questões objetivas. Para AF as discursivas não são usadas."""),
    ]

    secao_atual = ""
    for secao, pergunta, resposta in perguntas:
        if secao != secao_atual:
            st.subheader(secao)
            secao_atual = secao
        with st.expander(pergunta):
            st.markdown(resposta)


def main():
    st.set_page_config(
        page_title="Gerador de Provas Unicid",
        page_icon="📝",
        layout="wide"
    )

    st.title("Gerador de Provas Unicid")
    st.caption("Versão 3.6.3 · Prof.Me. Cid R. Andrade · profandrade@gmail.com · Co-Autor: Prof.Me. Rafael Cotrin (v3.4.0+)")

    tab_ger, tab_man, tab_faq = st.tabs(["🏠 Gerador de Provas", "📖 Manual", "❓ FAQ"])

    with tab_ger:
        st.divider()

        # Seção de importação via IA (opcional, independente do fluxo principal)
        _ui_importar_ia()

        st.divider()

    col_esq, col_dir = st.columns([1, 1], gap="large")

    with col_esq:
        # 1. Tipo de avaliação
        tipo_label = st.radio(
            "Tipo de avaliação",
            ["Regimental — AR", "Final — AF  (20 objetivas)"],
            horizontal=True
        )
        tipo_codigo           = "R" if "Regimental" in tipo_label else "F"
        qt_questoes_objetivas = 8 if tipo_codigo == "R" else 20

        # 2. Tipo de gabarito
        gabarito_tipo = st.radio(
            "Tipo de gabarito",
            ["Padrão", "Zipgrade"],
            horizontal=True
        )
        modelo_caminho = os.path.join(BASE_DIR, MODELOS[(tipo_codigo, gabarito_tipo)])

        st.divider()

        # 3. Número de versões
        qt_versoes = st.slider("Número de versões da prova", min_value=1, max_value=8, value=4)
        st.caption(f"Versões: {', '.join(LETRAS_PROVA[:qt_versoes])}")

        st.divider()

        # 4. Dados do professor e disciplina
        nome_professor  = st.text_input("Professor",  placeholder="Nome completo do professor")
        nome_disciplina = st.text_input("Disciplina", placeholder="Nome da disciplina")

        st.divider()

        # 5. Pontuação por seção (opcional — aparece no cabeçalho da prova)
        st.subheader("Pontuação")
        st.caption("Aparece no cabeçalho de cada seção da prova. Deixe em branco para omitir.")
        if tipo_codigo == "R":
            qt_dissertativas = st.radio(
                "Questões discursivas",
                [2, 3],
                horizontal=True,
                key="qt_dis_radio",
            )
            _val_dis = 3.0 / qt_dissertativas
            _por_questao_fmt = f"{_val_dis:.2f}".replace(".", ",") + " pt"
            pontos_dis = "3,00 pts"   # total da seção — vai para o cabeçalho do DOCX
            st.caption(
                f"Valor por questão discursiva: **{_por_questao_fmt}** · "
                f"Total discursivas: **3,00 pts** · "
                f"Total objetivas: **2,00 pts** (8 × 0,25 pt)"
            )
        else:
            qt_dissertativas = 0
            pontos_dis = ""
        pontos_obj = "2,00 pts" if tipo_codigo == "R" else ""

        st.divider()

        # 6. Espaço dissertativo (somente AR)
        linhas_por_questao = 8
        if tipo_codigo == "R":
            st.subheader("Espaço de resposta discursiva")
            st.caption(
                "As discursivas ocupam página inteira (1 coluna), iniciando sempre em "
                "página nova. As linhas abaixo são inseridas diretamente no DOCX."
            )
            linhas_por_questao = st.slider(
                "Linhas de resposta por questão",
                min_value=4, max_value=20, value=8
            )

        st.divider()

        # 6. Opção de PDF
        gerar_pdf = st.checkbox(
            "Incluir PDF no download",
            value=False,
            help=(
                "Converte cada DOCX para PDF via Microsoft Word. "
                "Feche o Word antes de gerar."
                if PDF_DISPONIVEL
                else "Indisponível — instale: pip install docx2pdf (requer Microsoft Word)"
            ),
            disabled=not PDF_DISPONIVEL
        )
        if not PDF_DISPONIVEL:
            st.caption("Para habilitar PDF: `pip install docx2pdf` (requer Microsoft Word)")

    with col_dir:
        _usando_ia = st.session_state.get("ia_confirmadas", False)
        _ia_obj    = st.session_state.get("ia_obj", [])
        _ia_dis    = st.session_state.get("ia_dis", [])

        # 7. Upload de questões objetivas
        st.subheader("Questões objetivas")
        if _usando_ia and _ia_obj:
            st.info(f"{len(_ia_obj)} questão(ões) objetiva(s) carregada(s) via IA.")
        st.caption(
            "Planilha XLSX: **coluna A** = enunciado · **coluna B** = resposta correta "
            "· **colunas C–F** = distratores.  \nPrimeira linha = primeira questão (sem cabeçalho).  \n"
            "Formatação: `*itálico*` · `**negrito**` · `***negrito e itálico***`"
            + ("  \n_Upload opcional — questões da IA estão ativas._" if (_usando_ia and _ia_obj) else "")
        )
        arquivo_objetivas = st.file_uploader(
            "Carregar questões objetivas (.xlsx)",
            type=["xlsx"],
            key="objetivas"
        )

        # 8. Upload de questões dissertativas (somente AR)
        arquivo_dissertativas = None
        if tipo_codigo == "R":
            st.subheader("Questões discursivas")
            if _usando_ia and _ia_dis:
                st.info(f"{len(_ia_dis)} questão(ões) discursiva(s) carregada(s) via IA.")
            st.caption(
                "Planilha XLSX: **coluna A** = enunciado.  \n"
                "Primeira linha = primeira questão (sem cabeçalho).  \n"
                "Formatação: `*itálico*` · `**negrito**` · `***negrito e itálico***`"
                + ("  \n_Upload opcional — questões da IA estão ativas._" if (_usando_ia and _ia_dis) else "")
            )
            arquivo_dissertativas = st.file_uploader(
                "Carregar questões discursivas (.xlsx)",
                type=["xlsx"],
                key="dissertativas"
            )

    st.divider()

    # 9. Imagens das questões (opcional)
    arquivos_imagens = st.file_uploader(
        "Imagens das questões (opcional)",
        type=["png", "jpg", "jpeg", "gif", "bmp"],
        accept_multiple_files=True,
        key="imagens_questoes",
        help=(
            "Para incluir uma imagem em uma questão, você tem três opções:  \n\n"
            "**1. Tag inline no XLSX (mais simples):** escreva `[img:nome.png]` "
            "em qualquer ponto do enunciado na coluna A. "
            "A tag é removida do texto da prova e a imagem aparece logo abaixo da questão.  \n\n"
            "**2. Coluna dedicada no XLSX:** informe só o nome do arquivo — "
            "coluna G para objetivas, coluna B para discursivas.  \n\n"
            "**3. DALL-E 3 (geração automática, sem upload):** abra o painel "
            "\"Importar documento não estruturado via IA\" (acima desta seção) "
            "e use a subseção \"Imagens com DALL-E 3\" para gerar imagens "
            "ilustrativas automaticamente a partir dos enunciados."
        ),
    )

    st.divider()

    # 10. Botão de geração
    gerar = st.button("Gerar Provas", type="primary", use_container_width=True)

    if gerar:
        _usando_ia = st.session_state.get("ia_confirmadas", False)
        _ia_obj    = st.session_state.get("ia_obj", [])
        _ia_dis    = st.session_state.get("ia_dis", [])

        # --- Validações ---
        if not nome_professor.strip():
            st.error("Preencha o nome do professor antes de continuar.")
            return
        if not nome_disciplina.strip():
            st.error("Preencha o nome da disciplina antes de continuar.")
            return

        tem_obj = arquivo_objetivas is not None or (_usando_ia and _ia_obj)
        if not tem_obj:
            st.error(
                "Carregue o arquivo de questões objetivas ou importe via IA antes de continuar."
            )
            return

        tem_dis = (
            arquivo_dissertativas is not None
            or (_usando_ia and _ia_dis)
        )
        if tipo_codigo == "R" and not tem_dis:
            st.error(
                "Carregue o arquivo de questões discursivas ou importe via IA antes de continuar."
            )
            return

        if not os.path.exists(modelo_caminho):
            st.error(
                f"Arquivo de modelo '{os.path.basename(modelo_caminho)}' não encontrado. "
                "Contate o administrador do sistema."
            )
            return

        # --- Geração ---
        with st.spinner("Lendo questões e gerando provas..."):

            # Monta dict de imagens carregadas: {filename: bytes}
            uploaded_images = {}
            if arquivos_imagens:
                for f_img in arquivos_imagens:
                    uploaded_images[f_img.name] = f_img.getvalue()

            # Fonte das objetivas: XLSX tem prioridade; IA como fallback
            imagens_obj = {}
            if arquivo_objetivas is not None:
                questoes_objetivas, imagens_obj = get_questoes_xlsx(arquivo_objetivas)
                if not questoes_objetivas:
                    st.error("Nenhuma questão objetiva válida encontrada. Verifique o arquivo XLSX.")
                    return
            else:
                questoes_objetivas = _ia_obj

            if len(questoes_objetivas) < qt_questoes_objetivas:
                st.error(
                    f"São necessárias **{qt_questoes_objetivas}** questões objetivas, "
                    f"mas foram encontradas apenas **{len(questoes_objetivas)}**."
                )
                return

            questoes_dissertativas   = None
            gabaritos_dissertativas  = None   # dict enunciado→resposta (só via IA)
            imagens_dis = {}
            # qt_dissertativas vem da UI (radio 2 ou 3), já definido acima
            if tipo_codigo == "R":
                if arquivo_dissertativas is not None:
                    questoes_dissertativas, imagens_dis = get_questoes_dissertativas_xlsx(arquivo_dissertativas)
                    if not questoes_dissertativas:
                        st.error("Nenhuma questão discursiva válida. Verifique o arquivo XLSX.")
                        return
                else:
                    questoes_dissertativas  = _ia_dis
                    _ia_dis_gab = st.session_state.get("ia_dis_gab", [])
                    # Monta mapeamento enunciado → resposta esperada
                    gabaritos_dissertativas = {
                        enunc: resp
                        for enunc, resp in zip(questoes_dissertativas, _ia_dis_gab)
                        if resp
                    }

                if len(questoes_dissertativas) < qt_dissertativas:
                    st.error(
                        f"São necessárias ao menos **{qt_dissertativas}** questões discursivas, "
                        f"mas foram encontradas apenas **{len(questoes_dissertativas)}**."
                    )
                    return

            gabaritos                     = {}
            gabaritos_dis_por_prova       = {}   # {nome_prova: [(enunciado, resposta)]}
            zip_buffer = BytesIO()
            erros      = []
            erros_pdf  = []

            # Monta dict {enunciado: bytes} — combina XLSX + IA (XLSX tem prioridade)
            imagens_mapa = {**imagens_obj, **imagens_dis}
            if imagens_mapa and uploaded_images is not None:
                faltando = [fname for fname in imagens_mapa.values() if fname not in uploaded_images]
                if faltando:
                    st.warning(
                        f"Imagem(ns) referenciada(s) mas não encontrada(s): "
                        f"**{', '.join(faltando)}**. "
                        "Faça o upload ou remova a tag `[img:]` da questão. "
                        "A prova será gerada sem essas imagens."
                    )
            imagens_xlsx = {
                enunc: uploaded_images[fname]
                for enunc, fname in imagens_mapa.items()
                if fname in uploaded_images
            } if uploaded_images else {}
            imagens_ia = st.session_state.get("ia_imgs", {})
            imagens_bytes = {**imagens_ia, **imagens_xlsx}  # XLSX sobrescreve IA

            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                for i in range(qt_versoes):
                    nome_prova = LETRAS_PROVA[i]
                    simbolo    = SIMBOLOS_PROVA[nome_prova]

                    questoes_selecionadas = random.sample(questoes_objetivas, qt_questoes_objetivas)
                    questoes_formatadas, gabarito_da_prova, enunciados_obj = criar_prova(
                        nome_prova, simbolo, qt_questoes_objetivas, questoes_selecionadas
                    )

                    dissertativas_selecionadas = None
                    if tipo_codigo == "R" and questoes_dissertativas:
                        dissertativas_selecionadas = random.sample(questoes_dissertativas, qt_dissertativas)
                        # Registra gabarito das dissertativas desta versão
                        if gabaritos_dissertativas:
                            gabaritos_dis_por_prova[nome_prova] = [
                                (enunc, gabaritos_dissertativas.get(enunc, ""))
                                for enunc in dissertativas_selecionadas
                            ]

                    gabaritos[nome_prova] = gabarito_da_prova

                    prova_bytes = gera_prova_bytes(
                        modelo_caminho, nome_prova, questoes_formatadas,
                        simbolo, tipo_codigo, dissertativas_selecionadas,
                        professor=nome_professor.strip(),
                        disciplina=nome_disciplina.strip(),
                        linhas_por_questao=linhas_por_questao,
                        pontos_obj=pontos_obj.strip(),
                        pontos_dis=pontos_dis,
                        imagens_questoes=imagens_bytes or None,
                        enunciados_obj=enunciados_obj,
                    )

                    if prova_bytes:
                        zf.writestr(f"prova_{nome_prova}.docx", prova_bytes)
                        if gerar_pdf:
                            pdf_bytes, pdf_erro = docx_bytes_to_pdf_bytes(prova_bytes)
                            if pdf_bytes:
                                zf.writestr(f"prova_{nome_prova}.pdf", pdf_bytes)
                            else:
                                erros_pdf.append((nome_prova, pdf_erro))
                    else:
                        erros.append(nome_prova)

                # Gabarito geral (com respostas dissertativas quando disponíveis)
                gabarito_str = gera_gabarito_str(
                    gabaritos, tipo_codigo,
                    gabaritos_dissertativas_por_prova=gabaritos_dis_por_prova or None,
                )
                data_hora        = datetime.now().strftime("%Y%m%d_%H%M%S")
                nome_completo_av = "Regimental" if tipo_codigo == "R" else "Final"
                nome_gabarito    = f"Gabarito_Geral_{nome_completo_av}_{data_hora}.txt"
                zf.writestr(nome_gabarito, gabarito_str.encode("utf-8"))

        if erros:
            st.warning(f"Erro ao gerar versões: {', '.join(erros)}. As demais foram incluídas no ZIP.")

        if erros_pdf:
            versoes_com_erro = ", ".join(v for v, _ in erros_pdf)
            # Mostra o motivo real do primeiro erro (geralmente o mesmo para todos)
            msg_erro = erros_pdf[0][1] or "erro desconhecido"
            st.warning(
                f"Conversão para PDF falhou nas versões: {versoes_com_erro}.  \n"
                f"**Motivo:** {msg_erro}  \n"
                "Dica: feche o Microsoft Word completamente antes de gerar."
            )

        if not erros:
            versoes_geradas = ", ".join(LETRAS_PROVA[:qt_versoes])
            extra = " + PDFs" if (gerar_pdf and PDF_DISPONIVEL and not erros_pdf) else ""
            st.success(f"{qt_versoes} versão(ões) gerada(s): {versoes_geradas}{extra}")

        zip_buffer.seek(0)
        data_hora = datetime.now().strftime("%Y%m%d_%H%M%S")
        nome_zip  = f"provas_{nome_completo_av}_{data_hora}.zip"

        st.download_button(
            label="Baixar ZIP com todas as provas",
            data=zip_buffer,
            file_name=nome_zip,
            mime="application/zip",
            use_container_width=True
        )

    with tab_man:
        _ui_manual()

    with tab_faq:
        _ui_faq()


if __name__ == "__main__":
    main()
